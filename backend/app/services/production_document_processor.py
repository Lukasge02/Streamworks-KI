"""
🏭 Production Document Processor for StreamWorks-KI
Advanced document processing with robust PDF, DOCX, and multi-format support

Author: Ravel-Lukas Geck
Company: Arvato Systems / Bertelsmann
Project: StreamWorks-KI Bachelor Thesis
Version: 2.0.0 (Production)
"""

import asyncio
import io
import logging
import mimetypes
import tempfile
import uuid
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any

from langchain.schema import Document
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    PythonCodeTextSplitter
)

from app.services.multi_format_processor import SupportedFormat, DocumentCategory

# Enhanced document loaders
try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        UnstructuredWordDocumentLoader,
        UnstructuredExcelLoader,
        UnstructuredHTMLLoader,
        CSVLoader,
        JSONLoader,
        TextLoader
    )
    PDF_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Some document loaders not available: {e}")
    PDF_AVAILABLE = False

# Additional processing libraries
try:
    import pypdf
    from docx import Document as DocxDocument
    import openpyxl
    from bs4 import BeautifulSoup
    ADVANCED_PROCESSING = True
except ImportError as e:
    logger.warning(f"Advanced processing libraries not available: {e}")
    ADVANCED_PROCESSING = False

logger = logging.getLogger(__name__)


class ProcessingQuality(Enum):
    """Quality levels for document processing"""
    EXCELLENT = "excellent"      # Perfect text extraction, all structure preserved
    GOOD = "good"               # Good text extraction, most structure preserved  
    ACCEPTABLE = "acceptable"   # Basic text extraction, some structure preserved
    POOR = "poor"              # Text extracted but structure lost
    FAILED = "failed"          # Processing failed, fallback used


@dataclass
class ProcessingResult:
    """Result of document processing"""
    success: bool
    documents: List[Document]
    file_format: SupportedFormat
    category: DocumentCategory
    processing_method: str
    chunk_count: int
    quality: ProcessingQuality
    extraction_confidence: float  # 0.0-1.0
    metadata: Dict[str, Any]
    error_message: Optional[str] = None
    warnings: List[str] = None
    processing_time: float = 0.0
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class ProductionDocumentProcessor:
    """
    🏭 Production-grade document processor with robust error handling
    
    Features:
    - Advanced PDF text extraction with fallback methods
    - DOCX processing with structure preservation
    - Excel sheet processing with table detection
    - HTML content extraction and cleaning
    - Intelligent chunking strategies per format
    - Quality assessment and confidence scoring
    - Comprehensive error handling and recovery
    """
    
    def __init__(self):
        self.supported_formats = {
            # Text formats
            '.txt': SupportedFormat.TXT,
            '.md': SupportedFormat.MD,
            '.markdown': SupportedFormat.MD,
            
            # PDF formats
            '.pdf': SupportedFormat.PDF,
            
            # Office formats
            '.docx': SupportedFormat.DOCX,
            '.doc': SupportedFormat.DOC,
            '.xlsx': SupportedFormat.XLSX,
            '.xls': SupportedFormat.XLS,
            
            # Web formats
            '.html': SupportedFormat.HTML,
            '.htm': SupportedFormat.HTM,
            
            # Data formats
            '.csv': SupportedFormat.CSV,
            '.json': SupportedFormat.JSON,
            '.jsonl': SupportedFormat.JSONL,
            '.xml': SupportedFormat.XML,
            
            # Code formats
            '.py': SupportedFormat.PY,
            '.js': SupportedFormat.JS,
            '.sql': SupportedFormat.SQL,
        }
        
        # Initialize text splitters
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
        )
        
        self.code_splitter = PythonCodeTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        logger.info("🏭 Production Document Processor initialized")
    
    async def process_document(
        self,
        file_path: str,
        content: bytes,
        filename: Optional[str] = None
    ) -> ProcessingResult:
        """
        Process document with production-grade error handling
        
        Args:
            file_path: Path to the file
            content: Raw file content as bytes
            filename: Optional filename override
            
        Returns:
            ProcessingResult with documents and metadata
        """
        start_time = asyncio.get_event_loop().time()
        processing_id = str(uuid.uuid4())[:8]
        
        try:
            # Detect file format
            file_format = self._detect_format(file_path, content)
            category = self._categorize_document(file_format, filename or file_path)
            
            logger.info(f"🏭 Processing document {processing_id}: {Path(file_path).name} ({file_format.value})")
            
            # Route to appropriate processor
            if file_format == SupportedFormat.PDF:
                result = await self._process_pdf_advanced(content, file_path)
            elif file_format in [SupportedFormat.DOCX, SupportedFormat.DOC]:
                result = await self._process_docx_advanced(content, file_path)
            elif file_format in [SupportedFormat.XLSX, SupportedFormat.XLS]:
                result = await self._process_excel_advanced(content, file_path)
            elif file_format in [SupportedFormat.HTML, SupportedFormat.HTM]:
                result = await self._process_html_advanced(content, file_path)
            elif file_format == SupportedFormat.CSV:
                result = await self._process_csv_advanced(content, file_path)
            elif file_format in [SupportedFormat.JSON, SupportedFormat.JSONL]:
                result = await self._process_json_advanced(content, file_path)
            elif file_format in [SupportedFormat.TXT, SupportedFormat.MD]:
                result = await self._process_text_advanced(content, file_path, file_format)
            elif file_format in [SupportedFormat.PY, SupportedFormat.JS, SupportedFormat.SQL]:
                result = await self._process_code_advanced(content, file_path, file_format)
            else:
                # Fallback to text processing
                result = await self._process_text_advanced(content, file_path, SupportedFormat.TXT)
            
            # Add processing metadata
            processing_time = asyncio.get_event_loop().time() - start_time
            
            return ProcessingResult(
                success=True,
                documents=result['documents'],
                file_format=file_format,
                category=category,
                processing_method=result['method'],
                chunk_count=len(result['documents']),
                quality=result['quality'],
                extraction_confidence=result['confidence'],
                metadata={
                    'processing_id': processing_id,
                    'file_size': len(content),
                    'original_filename': filename or Path(file_path).name,
                    **result.get('metadata', {})
                },
                warnings=result.get('warnings', []),
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = asyncio.get_event_loop().time() - start_time
            logger.error(f"❌ Document processing failed {processing_id}: {e}")
            
            # Emergency fallback
            try:
                fallback_docs = await self._emergency_fallback(content, file_path)
                return ProcessingResult(
                    success=False,
                    documents=fallback_docs,
                    file_format=SupportedFormat.TXT,
                    category=DocumentCategory.OFFICE_DOCUMENT,
                    processing_method="emergency_fallback",
                    chunk_count=len(fallback_docs),
                    quality=ProcessingQuality.FAILED,
                    extraction_confidence=0.1,
                    metadata={'processing_id': processing_id},
                    error_message=str(e),
                    processing_time=processing_time
                )
            except:
                return ProcessingResult(
                    success=False,
                    documents=[],
                    file_format=SupportedFormat.TXT,
                    category=DocumentCategory.OFFICE_DOCUMENT,
                    processing_method="complete_failure",
                    chunk_count=0,
                    quality=ProcessingQuality.FAILED,
                    extraction_confidence=0.0,
                    metadata={'processing_id': processing_id},
                    error_message=str(e),
                    processing_time=processing_time
                )
    
    async def _process_pdf_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced PDF processing with multiple extraction methods"""
        
        if not ADVANCED_PROCESSING:
            return await self._process_pdf_fallback(content, file_path)
        
        # Method 1: Try PyPDFLoader (most reliable)
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                loader = PyPDFLoader(tmp_file.name)
                documents = await asyncio.to_thread(loader.load)
                
                if documents and any(doc.page_content.strip() for doc in documents):
                    # Clean and enhance documents
                    enhanced_docs = []
                    for i, doc in enumerate(documents):
                        if doc.page_content.strip():
                            enhanced_doc = Document(
                                page_content=doc.page_content,
                                metadata={
                                    **doc.metadata,
                                    'source': file_path,
                                    'page_number': i + 1,
                                    'extraction_method': 'pypdf_loader',
                                    'chunk_type': 'page'
                                }
                            )
                            enhanced_docs.append(enhanced_doc)
                    
                    # Apply intelligent chunking
                    final_docs = self._apply_intelligent_chunking(enhanced_docs, SupportedFormat.PDF)
                    
                    return {
                        'documents': final_docs,
                        'method': 'pypdf_loader',
                        'quality': ProcessingQuality.EXCELLENT,
                        'confidence': 0.95,
                        'metadata': {
                            'total_pages': len(documents),
                            'text_pages': len(enhanced_docs)
                        }
                    }
        except Exception as e:
            logger.warning(f"PyPDFLoader failed: {e}")
        
        # Method 2: Try pypdf directly
        try:
            with io.BytesIO(content) as pdf_buffer:
                pdf_reader = pypdf.PdfReader(pdf_buffer)
                documents = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={
                                    'source': file_path,
                                    'page_number': page_num + 1,
                                    'extraction_method': 'pypdf_direct',
                                    'chunk_type': 'page'
                                }
                            )
                            documents.append(doc)
                    except Exception as page_error:
                        logger.warning(f"Failed to extract page {page_num}: {page_error}")
                
                if documents:
                    final_docs = self._apply_intelligent_chunking(documents, SupportedFormat.PDF)
                    return {
                        'documents': final_docs,
                        'method': 'pypdf_direct',
                        'quality': ProcessingQuality.GOOD,
                        'confidence': 0.85,
                        'metadata': {
                            'total_pages': len(pdf_reader.pages),
                            'extracted_pages': len(documents)
                        }
                    }
        except Exception as e:
            logger.warning(f"Direct pypdf extraction failed: {e}")
        
        # Method 3: Fallback to basic text extraction
        return await self._process_pdf_fallback(content, file_path)
    
    async def _process_pdf_fallback(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Fallback PDF processing when advanced methods fail"""
        try:
            # Try to extract any readable text
            text_content = content.decode('utf-8', errors='ignore')
            
            # Look for text patterns that might indicate actual content
            if any(pattern in text_content for pattern in ['stream', 'endstream', 'BT', 'ET']):
                # This looks like a PDF with text streams, but we can't extract it properly
                warning_text = (
                    f"⚠️ PDF processing limited - advanced text extraction not available.\n"
                    f"File: {Path(file_path).name}\n"
                    f"Size: {len(content)} bytes\n"
                    f"Status: Contains text but extraction failed\n\n"
                    f"To enable full PDF processing, ensure pypdf is properly installed.\n"
                    f"Raw content preview (first 500 chars):\n{text_content[:500]}..."
                )
            else:
                warning_text = f"PDF file detected but no readable text found: {Path(file_path).name}"
            
            doc = Document(
                page_content=warning_text,
                metadata={
                    'source': file_path,
                    'extraction_method': 'fallback_text',
                    'chunk_type': 'error_message',
                    'original_size': len(content)
                }
            )
            
            return {
                'documents': [doc],
                'method': 'fallback_text',
                'quality': ProcessingQuality.POOR,
                'confidence': 0.1,
                'warnings': ['PDF text extraction failed - using fallback method'],
                'metadata': {'fallback_reason': 'pypdf_not_available'}
            }
            
        except Exception as e:
            raise Exception(f"PDF fallback processing failed: {e}")
    
    async def _process_docx_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced DOCX processing with structure preservation and content cleaning"""
        
        if not ADVANCED_PROCESSING:
            return await self._process_text_advanced(content, file_path, SupportedFormat.TXT)
        
        try:
            # Method 1: Try python-docx for structured extraction
            with io.BytesIO(content) as docx_buffer:
                doc = DocxDocument(docx_buffer)
                
                documents = []
                current_section = ""
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Clean paragraph text
                        clean_text = self._clean_document_text(para.text)
                        if not clean_text:
                            continue
                            
                        # Detect headers and structure
                        if para.style.name.startswith('Heading'):
                            if current_section.strip():
                                # Save previous section
                                section_doc = Document(
                                    page_content=current_section.strip(),
                                    metadata={
                                        'source': file_path,
                                        'extraction_method': 'python_docx_cleaned',
                                        'chunk_type': 'section'
                                    }
                                )
                                documents.append(section_doc)
                            current_section = f"# {clean_text}\n\n"
                        else:
                            current_section += clean_text + "\n"
                
                # Add final section
                if current_section.strip():
                    section_doc = Document(
                        page_content=current_section.strip(),
                        metadata={
                            'source': file_path,
                            'extraction_method': 'python_docx_cleaned',
                            'chunk_type': 'section'
                        }
                    )
                    documents.append(section_doc)
                
                if documents:
                    final_docs = self._apply_intelligent_chunking(documents, SupportedFormat.DOCX)
                    return {
                        'documents': final_docs,
                        'method': 'python_docx_cleaned',
                        'quality': ProcessingQuality.EXCELLENT,
                        'confidence': 0.95,
                        'metadata': {
                            'total_paragraphs': len(doc.paragraphs),
                            'extracted_sections': len(documents),
                            'cleaning_applied': True
                        }
                    }
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
        
        # Method 2: Try UnstructuredWordDocumentLoader with aggressive cleaning
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                loader = UnstructuredWordDocumentLoader(tmp_file.name)
                documents = await asyncio.to_thread(loader.load)
                
                if documents:
                    enhanced_docs = []
                    all_content = ""
                    
                    # Combine all document content for aggressive cleaning
                    for doc in documents:
                        all_content += doc.page_content + "\n"
                    
                    # Apply aggressive cleaning to combined content
                    cleaned_content = self._clean_document_text_aggressive(all_content)
                    
                    if cleaned_content and len(cleaned_content) > 50:
                        # Create a single cleaned document
                        enhanced_doc = Document(
                            page_content=cleaned_content,
                            metadata={
                                'source': file_path,
                                'extraction_method': 'unstructured_aggressive_cleaned',
                                'cleaning_applied': True,
                                'original_length': len(all_content),
                                'cleaned_length': len(cleaned_content)
                            }
                        )
                        enhanced_docs.append(enhanced_doc)
                    
                    final_docs = self._apply_intelligent_chunking(enhanced_docs, SupportedFormat.DOCX)
                    return {
                        'documents': final_docs,
                        'method': 'unstructured_aggressive_cleaned',
                        'quality': ProcessingQuality.GOOD if enhanced_docs else ProcessingQuality.POOR,
                        'confidence': 0.85 if enhanced_docs else 0.3,
                        'metadata': {
                            'original_docs': len(documents),
                            'cleaned_docs': len(enhanced_docs),
                            'cleaning_applied': True,
                            'content_reduction': f"{len(all_content)} -> {len(cleaned_content) if cleaned_content else 0} chars"
                        }
                    }
        except Exception as e:
            logger.warning(f"UnstructuredWordDocumentLoader failed: {e}")
        
        # Fallback to text processing
        return await self._process_text_advanced(content, file_path, SupportedFormat.TXT)
    
    async def _process_excel_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced Excel processing with sheet and table detection"""
        
        if not ADVANCED_PROCESSING:
            return await self._process_text_advanced(content, file_path, SupportedFormat.TXT)
        
        try:
            with io.BytesIO(content) as excel_buffer:
                workbook = openpyxl.load_workbook(excel_buffer, data_only=True)
                documents = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    
                    # Extract sheet data
                    sheet_data = []
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                    
                    if sheet_data:
                        # Convert to readable format
                        sheet_content = f"## Sheet: {sheet_name}\n\n"
                        
                        # Check if first row looks like headers
                        if sheet_data and all(isinstance(cell, str) and cell.strip() for cell in sheet_data[0][:3]):
                            headers = sheet_data[0]
                            data_rows = sheet_data[1:]
                            
                            # Create table format
                            sheet_content += "| " + " | ".join(headers) + " |\n"
                            sheet_content += "|" + "|".join([" --- " for _ in headers]) + "|\n"
                            
                            for row in data_rows[:50]:  # Limit to first 50 rows
                                if any(cell.strip() for cell in row):
                                    sheet_content += "| " + " | ".join(row[:len(headers)]) + " |\n"
                        else:
                            # Simple format
                            for row in sheet_data[:50]:
                                if any(cell.strip() for cell in row):
                                    sheet_content += " | ".join(cell for cell in row if cell.strip()) + "\n"
                        
                        doc = Document(
                            page_content=sheet_content,
                            metadata={
                                'source': file_path,
                                'sheet_name': sheet_name,
                                'extraction_method': 'openpyxl',
                                'chunk_type': 'sheet',
                                'row_count': len(sheet_data)
                            }
                        )
                        documents.append(doc)
                
                if documents:
                    final_docs = self._apply_intelligent_chunking(documents, SupportedFormat.XLSX)
                    return {
                        'documents': final_docs,
                        'method': 'openpyxl',
                        'quality': ProcessingQuality.EXCELLENT,
                        'confidence': 0.9,
                        'metadata': {
                            'sheet_count': len(workbook.sheetnames),
                            'processed_sheets': len(documents)
                        }
                    }
        except Exception as e:
            logger.warning(f"openpyxl extraction failed: {e}")
        
        # Fallback to CSV-style processing
        try:
            text_content = content.decode('utf-8', errors='ignore')
            doc = Document(
                page_content=text_content,
                metadata={
                    'source': file_path,
                    'extraction_method': 'text_fallback',
                    'chunk_type': 'raw_text'
                }
            )
            return {
                'documents': [doc],
                'method': 'text_fallback',
                'quality': ProcessingQuality.POOR,
                'confidence': 0.3,
                'warnings': ['Excel processing failed - using text fallback']
            }
        except Exception as e:
            raise Exception(f"Excel processing completely failed: {e}")
    
    async def _process_html_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced HTML processing with content extraction"""
        
        try:
            html_content = content.decode('utf-8', errors='ignore')
            
            if ADVANCED_PROCESSING:
                # Use BeautifulSoup for clean text extraction
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text content
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                quality = ProcessingQuality.EXCELLENT
                confidence = 0.9
                method = 'beautifulsoup'
            else:
                # Basic HTML processing
                text = html_content
                quality = ProcessingQuality.ACCEPTABLE
                confidence = 0.6
                method = 'basic_html'
            
            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={
                        'source': file_path,
                        'extraction_method': method,
                        'chunk_type': 'html_content'
                    }
                )
                
                final_docs = self._apply_intelligent_chunking([doc], SupportedFormat.HTML)
                return {
                    'documents': final_docs,
                    'method': method,
                    'quality': quality,
                    'confidence': confidence
                }
            else:
                raise Exception("No text content found in HTML")
                
        except Exception as e:
            raise Exception(f"HTML processing failed: {e}")
    
    async def _process_csv_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced CSV processing with structure preservation"""
        
        try:
            text_content = content.decode('utf-8', errors='ignore')
            
            # Create temporary file for CSVLoader
            with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as tmp_file:
                tmp_file.write(text_content)
                tmp_file.flush()
                
                loader = CSVLoader(tmp_file.name)
                documents = await asyncio.to_thread(loader.load)
                
                if documents:
                    enhanced_docs = []
                    for i, doc in enumerate(documents):
                        enhanced_doc = Document(
                            page_content=doc.page_content,
                            metadata={
                                **doc.metadata,
                                'source': file_path,
                                'row_number': i + 1,
                                'extraction_method': 'csv_loader',
                                'chunk_type': 'csv_row'
                            }
                        )
                        enhanced_docs.append(enhanced_doc)
                    
                    final_docs = self._apply_intelligent_chunking(enhanced_docs, SupportedFormat.CSV)
                    return {
                        'documents': final_docs,
                        'method': 'csv_loader',
                        'quality': ProcessingQuality.EXCELLENT,
                        'confidence': 0.95,
                        'metadata': {'row_count': len(documents)}
                    }
        except Exception as e:
            logger.warning(f"CSV loader failed: {e}")
        
        # Fallback to text processing
        return await self._process_text_advanced(content, file_path, SupportedFormat.CSV)
    
    async def _process_json_advanced(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Advanced JSON processing with structure preservation"""
        
        try:
            text_content = content.decode('utf-8', errors='ignore')
            
            # Create temporary file for JSONLoader
            with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as tmp_file:
                tmp_file.write(text_content)
                tmp_file.flush()
                
                loader = JSONLoader(tmp_file.name, jq_schema='.', text_content=False)
                documents = await asyncio.to_thread(loader.load)
                
                if documents:
                    enhanced_docs = []
                    for doc in documents:
                        enhanced_doc = Document(
                            page_content=doc.page_content,
                            metadata={
                                **doc.metadata,
                                'source': file_path,
                                'extraction_method': 'json_loader',
                                'chunk_type': 'json_object'
                            }
                        )
                        enhanced_docs.append(enhanced_doc)
                    
                    final_docs = self._apply_intelligent_chunking(enhanced_docs, SupportedFormat.JSON)
                    return {
                        'documents': final_docs,
                        'method': 'json_loader',
                        'quality': ProcessingQuality.EXCELLENT,
                        'confidence': 0.95
                    }
        except Exception as e:
            logger.warning(f"JSON loader failed: {e}")
        
        # Fallback to text processing
        return await self._process_text_advanced(content, file_path, SupportedFormat.JSON)
    
    async def _process_text_advanced(self, content: bytes, file_path: str, file_format: SupportedFormat) -> Dict[str, Any]:
        """Advanced text processing with intelligent chunking"""
        
        try:
            text_content = content.decode('utf-8', errors='ignore')
            
            if not text_content.strip():
                raise Exception("Empty text content")
            
            doc = Document(
                page_content=text_content,
                metadata={
                    'source': file_path,
                    'extraction_method': 'text_decode',
                    'chunk_type': 'text_content',
                    'character_count': len(text_content)
                }
            )
            
            final_docs = self._apply_intelligent_chunking([doc], file_format)
            return {
                'documents': final_docs,
                'method': 'text_decode',
                'quality': ProcessingQuality.EXCELLENT,
                'confidence': 0.95
            }
            
        except Exception as e:
            raise Exception(f"Text processing failed: {e}")
    
    async def _process_code_advanced(self, content: bytes, file_path: str, file_format: SupportedFormat) -> Dict[str, Any]:
        """Advanced code processing with syntax-aware chunking"""
        
        try:
            code_content = content.decode('utf-8', errors='ignore')
            
            if not code_content.strip():
                raise Exception("Empty code content")
            
            doc = Document(
                page_content=code_content,
                metadata={
                    'source': file_path,
                    'extraction_method': 'code_decode',
                    'chunk_type': 'code_content',
                    'language': file_format.value
                }
            )
            
            # Use code-specific chunking
            chunks = self.code_splitter.split_documents([doc])
            
            enhanced_chunks = []
            for i, chunk in enumerate(chunks):
                enhanced_chunk = Document(
                    page_content=chunk.page_content,
                    metadata={
                        **chunk.metadata,
                        'chunk_index': i,
                        'chunk_type': 'code_chunk'
                    }
                )
                enhanced_chunks.append(enhanced_chunk)
            
            return {
                'documents': enhanced_chunks,
                'method': 'code_splitter',
                'quality': ProcessingQuality.EXCELLENT,
                'confidence': 0.95
            }
            
        except Exception as e:
            raise Exception(f"Code processing failed: {e}")
    
    async def _emergency_fallback(self, content: bytes, file_path: str) -> List[Document]:
        """Emergency fallback when all processing methods fail"""
        
        try:
            # Try UTF-8 decode
            text_content = content.decode('utf-8', errors='ignore')
        except:
            # Try latin-1 decode
            try:
                text_content = content.decode('latin-1', errors='ignore')
            except:
                # Last resort - create error message
                text_content = f"❌ Could not process file: {Path(file_path).name}\nSize: {len(content)} bytes"
        
        # Create minimal document
        doc = Document(
            page_content=text_content[:5000],  # Limit to first 5000 chars
            metadata={
                'source': file_path,
                'extraction_method': 'emergency_fallback',
                'chunk_type': 'fallback_content',
                'truncated': len(text_content) > 5000
            }
        )
        
        return [doc]
    
    def _apply_intelligent_chunking(self, documents: List[Document], file_format: SupportedFormat) -> List[Document]:
        """Apply intelligent chunking based on file format"""
        
        if not documents:
            return documents
        
        try:
            if file_format == SupportedFormat.MD:
                # Use markdown-aware chunking
                all_chunks = []
                for doc in documents:
                    chunks = self.markdown_splitter.split_text(doc.page_content)
                    for chunk in chunks:
                        if isinstance(chunk, Document):
                            chunk.metadata.update(doc.metadata)
                            all_chunks.append(chunk)
                        else:
                            # Handle simple text chunks
                            chunk_doc = Document(
                                page_content=chunk,
                                metadata={**doc.metadata, 'chunk_type': 'markdown_chunk'}
                            )
                            all_chunks.append(chunk_doc)
                return all_chunks if all_chunks else documents
            
            elif file_format in [SupportedFormat.PY, SupportedFormat.JS, SupportedFormat.SQL]:
                # Use code-aware chunking
                all_chunks = []
                for doc in documents:
                    chunks = self.code_splitter.split_documents([doc])
                    all_chunks.extend(chunks)
                return all_chunks if all_chunks else documents
            
            else:
                # Use general text chunking
                all_chunks = []
                for doc in documents:
                    chunks = self.text_splitter.split_documents([doc])
                    for i, chunk in enumerate(chunks):
                        chunk.metadata.update({
                            **doc.metadata,
                            'chunk_index': i,
                            'parent_doc_id': id(doc)
                        })
                    all_chunks.extend(chunks)
                return all_chunks if all_chunks else documents
        
        except Exception as e:
            logger.warning(f"Intelligent chunking failed: {e}, using original documents")
            return documents
    
    def _detect_format(self, file_path: str, content: bytes) -> SupportedFormat:
        """Detect file format from extension and content"""
        
        # Get extension
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in self.supported_formats:
            return self.supported_formats[file_ext]
        
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(file_path)
        
        mime_map = {
            'application/pdf': SupportedFormat.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': SupportedFormat.DOCX,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': SupportedFormat.XLSX,
            'text/html': SupportedFormat.HTML,
            'text/csv': SupportedFormat.CSV,
            'application/json': SupportedFormat.JSON,
            'text/plain': SupportedFormat.TXT,
        }
        
        if mime_type in mime_map:
            return mime_map[mime_type]
        
        # Content-based detection for common formats
        if content[:4] == b'%PDF':
            return SupportedFormat.PDF
        elif content[:2] == b'PK':  # ZIP-based formats (DOCX, XLSX)
            if b'word/' in content[:1000]:
                return SupportedFormat.DOCX
            elif b'xl/' in content[:1000]:
                return SupportedFormat.XLSX
        elif content.startswith(b'<html') or b'<HTML' in content[:100]:
            return SupportedFormat.HTML
        
        # Default to text
        return SupportedFormat.TXT
    
    def _clean_document_text(self, text: str) -> str:
        """Clean extracted text from HTML/XML markup and unwanted content"""
        import re
        
        if not text or not text.strip():
            return ""
        
        # Remove email headers and metadata
        email_patterns = [
            r'Date:\s*[^\n]+',
            r'Message-ID:\s*[^\n]+',
            r'Subject:\s*Exported From Confluence[^\n]*',
            r'MIME-Version:\s*[^\n]+',
            r'Content-Type:\s*[^\n]+',
            r'Content-Transfer-Encoding:\s*[^\n]+',
            r'Content-Location:\s*[^\n]+',
            r'------=_Part_\d+_\d+\.\d+',
        ]
        
        cleaned_text = text
        for pattern in email_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Remove HTML/XML tags and entities
        html_patterns = [
            r'<[^>]+>',  # HTML tags
            r'&[a-zA-Z0-9#]+;',  # HTML entities like &nbsp; &amp;
            r'=3D',  # Quoted-printable encoding
            r'=C3=',  # More quoted-printable
            r'=E2=80=',  # Unicode quoted-printable
            r'xmlns[^=]*=[^>\s]+',  # XML namespaces
            r'<!--[^>]*-->',  # HTML comments
            r'<!\[CDATA\[.*?\]\]>',  # CDATA sections
        ]
        
        for pattern in html_patterns:
            cleaned_text = re.sub(pattern, ' ', cleaned_text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove CSS styles and scripts
        css_script_patterns = [
            r'<style[^>]*>.*?</style>',
            r'<script[^>]*>.*?</script>',
            r'@media[^{]*\{[^}]*\}',
            r'\.[\w-]+\s*\{[^}]*\}',  # CSS classes
            r'#[\w-]+\s*\{[^}]*\}',   # CSS IDs
        ]
        
        for pattern in css_script_patterns:
            cleaned_text = re.sub(pattern, ' ', cleaned_text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove Office-specific markup
        office_patterns = [
            r'mso-[^:;]+:[^;]+;',  # Microsoft Office styles
            r'o:[\w-]+=[^>\s]+',   # Office namespaces
            r'w:[\w-]+=[^>\s]+',   # Word namespaces
            r'v:[\w-]+=[^>\s]+',   # VML namespaces
        ]
        
        for pattern in office_patterns:
            cleaned_text = re.sub(pattern, ' ', cleaned_text, flags=re.IGNORECASE)
        
        # Clean up whitespace and formatting
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Multiple spaces to single
        cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_text)  # Multiple newlines
        cleaned_text = re.sub(r'^\s+|\s+$', '', cleaned_text, flags=re.MULTILINE)  # Trim lines
        
        # Remove lines that are mostly markup or non-meaningful
        lines = cleaned_text.split('\n')
        meaningful_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip empty lines, lines with only symbols, or very short lines with mostly symbols
            if (len(line) > 5 and 
                not re.match(r'^[^\w\s]*$', line) and  # Not just symbols
                not re.match(r'^[0-9\s\-\.\:]+$', line) and  # Not just numbers/dates
                len(re.findall(r'[a-zA-ZäöüÄÖÜß]', line)) > 3):  # Has enough letters
                meaningful_lines.append(line)
        
        result = '\n'.join(meaningful_lines).strip()
        
        # Final cleanup: ensure we have meaningful content
        if len(result) < 20 or len(re.findall(r'[a-zA-ZäöüÄÖÜß]', result)) < 10:
            return ""
        
        return result
    
    def _clean_document_text_aggressive(self, text: str) -> str:
        """Aggressively clean document text for RAG-optimal content"""
        import re
        from bs4 import BeautifulSoup
        
        if not text or not text.strip():
            return ""
        
        logger.info(f"🧹 Aggressive cleaning started: {len(text)} chars")
        
        # Step 1: Use BeautifulSoup to strip HTML if available
        if ADVANCED_PROCESSING:
            try:
                soup = BeautifulSoup(text, 'html.parser')
                # Remove script and style elements completely
                for script in soup(["script", "style", "head", "meta", "link"]):
                    script.decompose()
                text = soup.get_text()
            except:
                pass  # Continue with regex cleaning
        
        # Step 2: Remove email headers and Confluence metadata  
        email_patterns = [
            r'Date:\s*[^\n]+',
            r'Message-ID:\s*[^\n]+', 
            r'Subject:\s*Exported From Confluence[^\n]*',
            r'MIME-Version:\s*[^\n]+',
            r'Content-Type:\s*[^\n]+',
            r'Content-Transfer-Encoding:\s*[^\n]+',
            r'Content-Location:\s*[^\n]+',
            r'------=_Part_\d+_\d+\.\d+[^\n]*',
            r'boundary="[^"]*"',
        ]
        
        for pattern in email_patterns:
            text = re.sub(pattern, '', text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Step 3: Remove HTML/XML completely
        html_patterns = [
            r'<[^>]*>',  # All HTML tags
            r'&[a-zA-Z0-9#]+;',  # HTML entities
            r'=3D[A-Fa-f0-9]*',  # Quoted-printable
            r'=C3=[A-Fa-f0-9]*',  # Unicode quoted-printable  
            r'=E2=80=[A-Fa-f0-9]*',  # More Unicode
            r'xmlns[^=]*=[^>\s]+',  # XML namespaces
            r'<!--.*?-->',  # Comments
            r'<!\[CDATA\[.*?\]\]>',  # CDATA
        ]
        
        for pattern in html_patterns:
            text = re.sub(pattern, ' ', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Step 4: Remove Office/CSS styling completely
        office_patterns = [
            r'mso-[^:;]+:[^;]+;',
            r'font-[^:;]+:[^;]+;',
            r'margin[^:;]*:[^;]+;',
            r'padding[^:;]*:[^;]+;',
            r'border[^:;]*:[^;]+;',
            r'style="[^"]*"',
            r'class="[^"]*"',
            r'data-[^=]*="[^"]*"',
            r'width="[^"]*"',
            r'height="[^"]*"',
        ]
        
        for pattern in office_patterns:
            text = re.sub(pattern, ' ', text, flags=re.IGNORECASE)
        
        # Step 5: Clean up structure and formatting
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces
        text = re.sub(r'\n\s*\n+', '\n\n', text)  # Multiple newlines
        
        # Step 6: Extract meaningful lines only
        lines = text.split('\n')
        meaningful_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip empty or very short lines
            if len(line) < 6:
                continue
                
            # Skip lines that are mostly symbols/markup
            if re.match(r'^[^\w\s]*$', line):
                continue
                
            # Skip lines with mostly numbers/dates/IDs
            if re.match(r'^[0-9\s\-\.\:\(\)]+$', line):
                continue
                
            # Skip CSS/style remnants
            if re.search(r'(color:|font-|margin|padding|border)', line, re.IGNORECASE):
                continue
            
            # Skip URLs and email addresses
            if re.search(r'(https?://|@.*\.com|\.html|\.css)', line):
                continue
                
            # Must have enough real letters (not just symbols)
            letter_count = len(re.findall(r'[a-zA-ZäöüÄÖÜßÀ-ÿ]', line))
            if letter_count < 5:
                continue
                
            # Skip lines that look like metadata
            if re.match(r'^(width|height|size|data-)', line, re.IGNORECASE):
                continue
                
            meaningful_lines.append(line)
        
        # Step 7: Reconstruct content with proper structure
        result = '\n'.join(meaningful_lines).strip()
        
        # Step 8: Final validation
        if len(result) < 50:
            logger.warning(f"⚠️ Aggressive cleaning resulted in very short content: {len(result)} chars")
            return ""
            
        letter_count = len(re.findall(r'[a-zA-ZäöüÄÖÜßÀ-ÿ]', result))
        if letter_count < 30:
            logger.warning(f"⚠️ Aggressive cleaning resulted in content with too few letters: {letter_count}")
            return ""
        
        logger.info(f"✅ Aggressive cleaning completed: {len(text)} -> {len(result)} chars ({len(meaningful_lines)} meaningful lines)")
        
        return result
    
    def _categorize_document(self, file_format: SupportedFormat, filename: str) -> DocumentCategory:
        """Categorize document based on format and filename"""
        
        filename_lower = filename.lower()
        
        # Category mapping based on content and filename
        if any(term in filename_lower for term in ['help', 'hilfe', 'anleitung', 'manual', 'guide', 'faq']):
            return DocumentCategory.HELP_DOCUMENTATION
        elif any(term in filename_lower for term in ['template', 'vorlage', 'schema', 'xml']):
            return DocumentCategory.XML_CONFIGURATION
        elif any(term in filename_lower for term in ['config', 'konfig', 'setting', 'einstellung']):
            return DocumentCategory.CONFIGURATION
        elif file_format in [SupportedFormat.PY, SupportedFormat.JS, SupportedFormat.SQL]:
            return DocumentCategory.CODE_SCRIPT
        elif file_format in [SupportedFormat.CSV, SupportedFormat.JSON, SupportedFormat.XLSX]:
            return DocumentCategory.STRUCTURED_DATA
        elif file_format in [SupportedFormat.PDF, SupportedFormat.DOCX]:
            return DocumentCategory.OFFICE_DOCUMENT
        else:
            return DocumentCategory.HELP_DOCUMENTATION


# Global instance
production_document_processor = ProductionDocumentProcessor()