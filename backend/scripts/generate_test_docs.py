#!/usr/bin/env python3
"""
Generiert 50 Testdokumente in 5 Kategorien und laedt sie via API hoch.

Voraussetzung: Backend + Qdrant + MinIO muessen laufen.
Aufruf: python backend/scripts/generate_test_docs.py
"""

import os
import sys
import csv
import io
import tempfile
import time
import httpx

# ── Konfiguration ────────────────────────────────────────────────────

BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8000")
API = f"{BACKEND_URL}/api/documents"
OUTPUT_DIR = os.path.join(tempfile.gettempdir(), "streamworks_testdocs")

FOLDER_COLORS = {
    "SAP-Dokumentation": "#0066cc",
    "Anforderungen": "#059669",
    "Betriebshandbuecher": "#d97706",
    "Konfigurationen": "#7c3aed",
    "Berichte": "#dc2626",
}

# ── Inhalts-Templates ────────────────────────────────────────────────

SAP_DOCS = [
    ("SAP_MM_Bestellprozess", "SAP MM Bestellprozess",
     """SAP MM Bestellprozess - Dokumentation

1. Uebersicht
Der Bestellprozess im SAP-Modul MM (Materials Management) umfasst die komplette Beschaffungskette von der Bedarfsermittlung bis zum Wareneingang. Streamworks automatisiert die naechtliche Verarbeitung der Bestellanforderungen (BANFen) und deren Umwandlung in Bestellungen.

2. Prozessablauf
2.1 Bedarfsermittlung
Die Bedarfsermittlung erfolgt ueber MRP (Material Requirements Planning) im SAP. Der Dispositionslauf (Transaktion MD01/MD02) erzeugt automatisch Bestellanforderungen fuer Materialien, deren Bestand unter den Meldebestand faellt.

2.2 Bestellanforderung (BANF)
BANFen werden ueber Transaktion ME51N manuell oder automatisch erstellt. Streamworks ueberwacht den Status der BANFen und loest bei Freigabe automatisch den naechsten Verarbeitungsschritt aus.

2.3 Bestellung
Die Umwandlung von BANFen in Bestellungen erfolgt ueber ME21N. Der Streamworks-Job SAP_MM_BANF_TO_BEST fuehrt diesen Schritt automatisiert durch und protokolliert die Ergebnisse.

2.4 Wareneingang
Der Wareneingang wird ueber MIGO gebucht. Streamworks ueberwacht eingehende IDocs und verarbeitet automatische Wareneingangsbuchungen.

3. Streamworks-Integration
Fuer den SAP MM Bestellprozess sind folgende Streams konfiguriert:
- GECK003_SAP_MM_BANF: Taegliche Verarbeitung offener BANFen
- GECK003_SAP_MM_BEST: Bestellerzeugung mit SAP-RFC-Aufruf
- GECK003_SAP_MM_WE: Wareneingangsverarbeitung

4. Fehlerbehandlung
Bei Fehlern im Bestellprozess wird automatisch eine Incident-Meldung erzeugt. Der Recovery-Mechanismus versucht fehlgeschlagene Jobs bis zu 3 Mal erneut auszufuehren. Die maximale Wartezeit zwischen Versuchen betraegt 15 Minuten.

5. Monitoring
Die Ueberwachung erfolgt ueber das Streamworks-Dashboard. Kritische Metriken sind:
- Anzahl verarbeiteter BANFen pro Tag
- Durchschnittliche Verarbeitungszeit pro Bestellung
- Fehlerquote bei RFC-Aufrufen
- Queue-Laenge der wartenden Auftraege"""),

    ("SAP_FI_Buchungslogik", "SAP FI Buchungslogik",
     """SAP FI Buchungslogik - Technische Dokumentation

1. Einleitung
Die Buchungslogik im SAP Financial Accounting (FI) bildet das Rueckgrat der finanziellen Datenverarbeitung. Dieses Dokument beschreibt die Integration von Streamworks mit den SAP FI Buchungsprozessen.

2. Buchungsarten
2.1 Sachkontenbuchungen
Sachkontenbuchungen (Transaktion FB50) werden ueber Batch-Input-Mappen verarbeitet. Streamworks steuert die naechtliche Verarbeitung dieser Mappen und ueberwacht den erfolgreichen Abschluss.

2.2 Kreditorenbuchungen
Eingangsrechnungen werden ueber FB60/MIRO erfasst. Der automatisierte Zahlungslauf (F110) wird durch Streamworks-Jobs zeitgesteuert gestartet.

2.3 Debitorenbuchungen
Ausgangsrechnungen und Mahnlaeufe (F150) werden ebenfalls ueber Streamworks automatisiert. Die Mahnparameter werden aus der Konfigurationstabelle gelesen.

3. Periodenabschluss
Der Monatsabschluss erfordert eine definierte Reihenfolge von SAP-Transaktionen:
- FAGL_FC_VALUATION: Fremdwaehrungsbewertung
- F101: Saldovortrag
- ASKB: Anlagenbuchhaltung Periodenabschluss

Streamworks stellt die korrekte Reihenfolge ueber Stream-Dependencies sicher. Ein Master-Stream GECK003_FI_MONATSABSCHLUSS koordiniert alle Teilschritte.

4. Schnittstellen
Die SAP FI Integration nutzt:
- RFC-Aufrufe fuer synchrone Verarbeitung
- IDoc-Schnittstelle fuer asynchrone Belegverarbeitung
- Batch-Input fuer Massenbuchungen

5. Protokollierung
Alle Buchungsvorgaenge werden im Streamworks Job-Log protokolliert. Die Aufbewahrungsfrist betraegt 90 Tage. Aeltere Logs werden automatisch archiviert."""),

    ("SAP_SD_Auftragsabwicklung", "SAP SD Auftragsabwicklung",
     """SAP SD Auftragsabwicklung - Prozessdokumentation

1. Prozessbeschreibung
Das SAP Sales and Distribution (SD) Modul verwaltet den gesamten Verkaufsprozess. Streamworks automatisiert die Hintergrundverarbeitung von Kundenauftraegen, Lieferungen und Fakturierung.

2. Auftragserfassung
Kundenauftraege werden ueber VA01 erfasst oder via EDI/IDoc automatisch angelegt. Streamworks ueberwacht die IDoc-Eingangsverarbeitung und loest bei neuen Auftraegen die Verfuegbarkeitspruefung aus.

3. Lieferungsabwicklung
Die Lieferungserstellung (VL01N) wird ueber Streamworks-Jobs gesteuert. Der Job GECK003_SD_LIEFERUNG laeuft taeglich um 06:00 Uhr und erstellt Lieferungen fuer alle faelligen Auftraege.

4. Fakturierung
Die Fakturierung (VF01) wird durch einen separaten Stream GECK003_SD_FAKTURA angestossen. Die Abhaengigkeit zum Lieferungs-Stream stellt sicher, dass nur fakturiert wird, wenn die Lieferung erfolgreich war.

5. Kreditmanagement
Die Kreditpruefung wird vor der Lieferungsfreigabe automatisch durchgefuehrt. Bei Kreditlimit-Ueberschreitung wird der Auftrag blockiert und eine Benachrichtigung an den Kreditmanager gesendet.

6. Reporting
Tagesberichte werden automatisch um 22:00 Uhr generiert:
- Auftragseingangsstatistik
- Lieferperformance
- Fakturauebersicht
- Kreditblock-Report"""),

    ("SAP_BC_Systemadministration", "SAP BC Systemadministration",
     """SAP Basis/BC Systemadministration - Betriebshandbuch

1. Systemlandschaft
Die SAP-Systemlandschaft umfasst drei Systeme (Entwicklung, Qualitaetssicherung, Produktion). Streamworks verwaltet systemuebergreifende Jobs und Transportwege.

2. Transportwesen
2.1 Transportauftraege
Streamworks ueberwacht den Transport von Entwicklungsobjekten durch die Systemlandschaft. Der Job GECK003_BC_TRANSPORT prueft stuendlich auf freigegebene Transportauftraege.

2.2 Importqueue
Die Importqueue des Produktivsystems wird ueber einen dedizierten Stream verwaltet. Importe werden nur im definierten Wartungsfenster (Sonntag 02:00-06:00) durchgefuehrt.

3. Hintergrund-Jobs
3.1 Job-Ueberwachung
Streamworks ueberwacht alle SAP-Hintergrund-Jobs ueber die XBP-Schnittstelle (External Background Processing). Fehlgeschlagene Jobs werden automatisch zur Wiederholung eingeplant.

3.2 Job-Scheduling
Die Einplanung von SAP-Jobs erfolgt zentral ueber Streamworks. Direkte Einplanung in SM36 ist fuer produktive Jobs nicht erlaubt.

4. Performance-Monitoring
- Taegliche Pruefung der Systemperformance (ST03N)
- Warnungen bei ABAP-Dumps (ST22)
- Spool-Bereinigung (SP01) alle 7 Tage
- Batch-Input-Protokoll-Loeschung nach 30 Tagen

5. Sicherheit
- Benutzersperren bei 5 Fehlversuchen
- Passwortaenderung alle 90 Tage
- Berechtigungspruefung ueber SU53-Auswertung
- Audit-Log-Auswertung woechentlich"""),

    ("SAP_HR_Personalabrechnung", "SAP HR Personalabrechnung",
     """SAP HR Personalabrechnung - Automatisierungsdokumentation

1. Abrechnungszyklus
Die monatliche Personalabrechnung wird vollstaendig ueber Streamworks automatisiert. Der Master-Stream GECK003_HR_ABRECHNUNG koordiniert alle Teilschritte.

2. Vorarbeiten
Vor der Abrechnung werden folgende Pruefungen durchgefuehrt:
- Zeitdaten-Vollstaendigkeitspruefung (PT60)
- Infotyp-Konsistenzpruefung
- Abrechnungskreis-Sperre setzen (PA03)

3. Abrechnungslauf
Der eigentliche Abrechnungslauf (PC00_M01_CALC) wird in mehreren Schritten ausgefuehrt:
- Simulation (Trockenlauf ohne Buchung)
- Ergebnispruefung und Freigabe
- Produktiver Abrechnungslauf
- Bankueberweisungsdatei erstellen (PC00_M01_CDTA)

4. Nacharbeiten
- Buchungsbeleg erzeugen (PC00_M01_HRF)
- Lohnsteueranmeldung erstellen
- Sozialversicherungsmeldungen generieren
- Lohnabrechnungen drucken/versenden

5. Fehlerbehandlung
Bei Abrechnungsfehlern werden die betroffenen Personalnummern in eine Wiedervorlage-Liste geschrieben. Der zustaendige Sachbearbeiter erhaelt eine automatische Benachrichtigung per E-Mail.

6. Archivierung
Abrechnungsergebnisse werden fuer 10 Jahre archiviert. Die Archivierung erfolgt quartalsweise ueber den Streamworks-Job GECK003_HR_ARCHIV."""),

    ("SAP_PP_Fertigungssteuerung", "SAP PP Fertigungssteuerung",
     """SAP PP Fertigungssteuerung - Integrationsdokumentation

1. Fertigungsplanung
Die Fertigungsplanung im SAP PP (Production Planning) wird ueber Streamworks mit der Materialwirtschaft und dem Vertrieb synchronisiert.

2. Planungslaeufe
2.1 MRP-Lauf
Der naechtliche MRP-Lauf (MD01) wird um 01:00 Uhr gestartet. Streamworks stellt sicher, dass vorher alle Bestandsbuchungen des Vortages abgeschlossen sind.

2.2 Kapazitaetsplanung
Die Kapazitaetsplanung (CM01-CM07) wird nach dem MRP-Lauf durchgefuehrt. Bei Kapazitaetsengpaessen werden automatisch Benachrichtigungen an die Fertigungsleitung versendet.

3. Fertigungsauftraege
- Planauftragsumsetzung (CO40N): Taeglich um 04:00 Uhr
- Rueckmeldungen verarbeiten (CO11N): Alle 2 Stunden
- Warenausgabe (MIGO): Bei Fertigungsauftragfreigabe

4. Betriebsdatenerfassung
BDE-Terminals melden Fertigungsfortschritte an SAP. Streamworks verarbeitet die eingehenden Rueckmeldungen und aktualisiert den Fertigungsstatus.

5. Qualitaetssicherung
Qualitaetsprueflose werden automatisch bei Wareneingang und Fertigmeldung erzeugt. Der Stream GECK003_PP_QM ueberwacht offene Prueflose und eskaliert bei Fristenueberschreitung."""),

    ("SAP_WM_Lagerverwaltung", "SAP WM Lagerverwaltung",
     """SAP WM/EWM Lagerverwaltung - Streamworks-Integration

1. Lagerprozesse
Die Lagerverwaltung nutzt SAP WM (Warehouse Management) fuer die Steuerung aller Lagerbewegungen. Streamworks automatisiert die Hintergrundverarbeitung.

2. Wareneingangsprozess
2.1 Avisierung
Eingehende Lieferungen werden ueber EDI avisiert. Streamworks verarbeitet die Avis-IDocs und erzeugt Transportbedarfe im Lager.

2.2 Einlagerung
Die Einlagerungsstrategie wird systemseitig ermittelt. Der Job GECK003_WM_EINLAGER erstellt Transportauftraege fuer die optimale Lagerplatzfindung.

3. Kommissionierung
Kommissionierauftraege werden aus Lieferungen abgeleitet. Streamworks gruppiert Auftraege nach Lagerbereich und optimiert die Kommissionierreihenfolge.

4. Inventur
- Stichtagsinventur: Jaehrlich zum 31.12., gesteuert ueber Master-Stream
- Permanente Inventur: Taeglich fuer ABC-Artikel
- Inventurdifferenzen werden automatisch gebucht und gemeldet

5. Bestandsoptimierung
Woechentliche Analysen zur Lageroptimierung:
- Lagerhuetkurvenanalyse
- ABC-XYZ-Klassifizierung aktualisieren
- Mindestbestaende pruefen und anpassen"""),

    ("SAP_BW_Reporting", "SAP BW Reporting und Datenextraktion",
     """SAP BW Reporting - Datenextraktions-Dokumentation

1. Datenextraktion
Streamworks steuert die naechtliche Datenextraktion aus den operativen SAP-Systemen in das Business Warehouse (BW).

2. Extraktionsprozesse
2.1 Full-Load
Vollstaendige Datenladungen werden bei Erstbefuellung und nach Datenbereinigungen durchgefuehrt. Der Master-Stream GECK003_BW_FULLLOAD koordiniert alle DataSource-Extraktionen.

2.2 Delta-Load
Taeglich werden nur geaenderte Datensaetze extrahiert. Die Delta-Queue wird um 23:00 Uhr verarbeitet. Streamworks ueberwacht den Delta-Status aller InfoProvider.

3. Prozessketten
BW-Prozessketten werden ueber Streamworks gestartet und ueberwacht:
- Extraktion -> Transformation -> Laden -> Aktivierung
- Bei Fehler in einem Schritt wird die gesamte Kette angehalten

4. Berichtsgenerierung
Automatische Berichtserstellung nach erfolgreicher Datenladung:
- Management-Dashboard aktualisieren
- PDF-Reports generieren und per E-Mail versenden
- Excel-Exporte in SharePoint ablegen

5. Datenqualitaet
Automatische Datenpruefungen nach jedem Ladelauf:
- Vollstaendigkeitspruefung (Satzanzahl-Vergleich)
- Plausibilitaetspruefung (Summen, Zeitreihen)
- Warnmeldung bei signifikanten Abweichungen"""),

    ("SAP_PI_Schnittstellenmanagement", "SAP PI/PO Schnittstellenmanagement",
     """SAP PI/PO Schnittstellenmanagement - Betriebsdokumentation

1. Architektur
SAP Process Integration (PI) / Process Orchestration (PO) bildet den zentralen Integrationsknoten. Streamworks koordiniert die Schnittstellenverarbeitung.

2. Schnittstellentypen
2.1 Synchrone Schnittstellen
- RFC-basierte Echtzeit-Kommunikation
- Web-Service-Aufrufe (SOAP/REST)
- Streamworks ueberwacht Antwortzeiten und Verfuegbarkeit

2.2 Asynchrone Schnittstellen
- IDoc-basierte Nachrichtenverarbeitung
- File-basierte Schnittstellen (CSV, XML, Flatfile)
- Streamworks steuert Polling-Intervalle und Fehlerbehandlung

3. Monitoring
Der Stream GECK003_PI_MONITOR laeuft alle 15 Minuten:
- Pruefung des Message-Status in PI
- Identifikation fehlerhafter Nachrichten
- Automatischer Restart bei bekannten Fehlern
- Eskalation bei unbekannten Fehlern

4. Alerting
Schwellwerte fuer automatische Alarme:
- Queue-Tiefe > 1000 Nachrichten
- Fehlerquote > 5%
- Antwortzeit > 30 Sekunden
- Kein Heartbeat seit > 5 Minuten

5. Wartung
- Cache-Bereinigung woechentlich
- Log-Archivierung nach 60 Tagen
- Channel-Restart bei Speicherueberlauf"""),

    ("SAP_SolMan_Changemanagement", "SAP Solution Manager Changemanagement",
     """SAP Solution Manager - Change Management Prozess

1. Change-Request-Workflow
Aenderungen an der SAP-Systemlandschaft werden ueber den Solution Manager gesteuert. Streamworks automatisiert Teile des Change-Management-Prozesses.

2. Normalchange
2.1 Antragstellung
Changes werden im Solution Manager als Service-Request erfasst. Streamworks prueft stuendlich auf neue freigegebene Change-Requests.

2.2 Transportsteuerung
Bei genehmigten Changes steuert Streamworks automatisch:
- Transport-Freigabe in Entwicklung
- Import in Qualitaetssystem
- Testdurchfuehrung anstosse
- Import in Produktion (nach Freigabe)

3. Emergency Change
Notfall-Changes werden mit erhoehter Prioritaet behandelt:
- Sofortiger Import ohne Wartungsfenster
- Automatische Benachrichtigung aller Stakeholder
- Nachtraegliche Dokumentation wird eingefordert

4. Reporting
- Woechentlicher Change-Report
- KPI-Dashboard (Durchlaufzeit, Erfolgsquote, Rollbacks)
- Compliance-Pruefung (ITIL-Konformitaet)

5. Integration
Der Solution Manager ist ueber RFC mit allen SAP-Systemen verbunden. Streamworks nutzt die ChaRM-API fuer die Statusaktualisierung von Transport-Requests."""),
]

ANFORDERUNGEN_DOCS = [
    ("Anforderung_Dateitransfer", "Fachliche Anforderung: Automatisierter Dateitransfer",
     """Fachliche Anforderung: Automatisierter Dateitransfer zwischen Standorten

Projekt: Streamworks Erweiterung
Version: 1.2
Autor: IT-Abteilung
Datum: 2025-11-15

1. Ausgangssituation
Aktuell werden Dateien zwischen den Standorten Hamburg, Muenchen und Berlin manuell per SFTP uebertragen. Dies ist fehleranfaellig und zeitaufwaendig.

2. Anforderung
Es soll ein automatisierter Dateitransfer-Stream in Streamworks eingerichtet werden, der:
- Taeglich um 22:00 Uhr alle Exportdateien sammelt
- Dateien komprimiert und verschluesselt
- An alle Zielstandorte verteilt
- Uebertragungsprotokoll erstellt
- Bei Fehler automatisch 3 Wiederholungsversuche durchfuehrt

3. Technische Rahmenbedingungen
- Protokoll: SFTP mit Zertifikatsauthentifizierung
- Maximale Dateigroesse: 500 MB
- Komprimierung: GZIP
- Verschluesselung: AES-256

4. Akzeptanzkriterien
- Alle Dateien werden innerhalb von 2 Stunden uebertragen
- Fehlerquote < 0.1%
- Lueckenlose Protokollierung
- Automatische Eskalation bei 3 aufeinanderfolgenden Fehlern"""),

    ("Anforderung_Monitoring", "Fachliche Anforderung: Job-Monitoring Dashboard",
     """Fachliche Anforderung: Erweitertes Job-Monitoring Dashboard

Projekt: Streamworks UI Erweiterung
Version: 2.0
Autor: Operations Team

1. Ist-Zustand
Das aktuelle Monitoring zeigt nur den letzten Status eines Jobs. Historische Daten und Trends sind nicht einsehbar.

2. Soll-Zustand
Ein erweitertes Dashboard soll folgende Informationen bereitstellen:
- Echtzeit-Status aller aktiven Streams
- Historische Ausfuehrungszeiten (Trend-Diagramm)
- Fehlerstatistiken pro Stream/Job
- SLA-Ueberwachung mit Ampelsystem
- Kapazitaetsauslastung der Agenten

3. Nicht-funktionale Anforderungen
- Aktualisierungsintervall: maximal 30 Sekunden
- Unterstuetzung von Filtern nach Agent, Stream, Zeitraum
- Export als PDF/Excel
- Responsive Design fuer Tablet-Nutzung

4. Priorisierung
Muss: Echtzeit-Status, Fehlerstatistik
Soll: Trend-Diagramme, SLA-Ampel
Kann: Kapazitaetsauslastung, PDF-Export"""),

    ("Anforderung_Benachrichtigung", "Fachliche Anforderung: Benachrichtigungssystem",
     """Fachliche Anforderung: Intelligentes Benachrichtigungssystem

1. Motivation
Aktuell werden alle Job-Fehler gleichermassen per E-Mail gemeldet. Dies fuehrt zu Alert-Fatigue und verzoegerter Reaktion bei kritischen Fehlern.

2. Anforderungen
2.1 Priorisierung
- Kritisch: Produktionsstopp, Datenverlust -> SMS + E-Mail + Teams-Nachricht
- Hoch: SLA-Verletzung, mehrfache Fehler -> E-Mail + Teams-Nachricht
- Mittel: Einzelner Fehlschlag mit Auto-Recovery -> E-Mail
- Niedrig: Warnungen, Performance-Degradation -> Dashboard-Eintrag

2.2 Eskalation
Wenn ein kritischer Alarm nach 15 Minuten nicht quittiert wird, erfolgt automatische Eskalation an den naechsthoehere Bereitschaftsebene.

2.3 Aggregation
Gleiche Fehlermeldungen werden innerhalb von 5 Minuten aggregiert. Statt 50 einzelner Mails wird eine zusammenfassende Nachricht versendet.

3. Kanaele
- E-Mail (SMTP)
- Microsoft Teams (Webhook)
- SMS (ueber Gateway-API)
- Streamworks-Dashboard (WebSocket)"""),

    ("Anforderung_Versionierung", "Fachliche Anforderung: Stream-Versionierung",
     """Fachliche Anforderung: Stream-Versionierung und Rollback

1. Problem
Aenderungen an Stream-Definitionen koennen nicht zurueckverfolgt werden. Bei fehlerhaften Aenderungen gibt es keinen einfachen Rollback-Mechanismus.

2. Loesung
Implementierung eines Versionierungssystems fuer Stream-Definitionen:
- Jede Aenderung erzeugt eine neue Version
- Versionshistorie mit Diff-Ansicht
- Ein-Klick-Rollback auf beliebige Vorgaengerversion
- Vergleich von zwei Versionen nebeneinander

3. Technische Umsetzung
- Speicherung in Git-aehnlicher Struktur
- Versionsnummer: MAJOR.MINOR (z.B. 1.0, 1.1, 2.0)
- Major-Version bei strukturellen Aenderungen
- Minor-Version bei Parameteraenderungen

4. Berechtigungen
- Erstellen neuer Versionen: Stream-Owner, Admins
- Rollback: Nur Admins, erfordert 4-Augen-Prinzip
- Ansicht Versionshistorie: Alle authentifizierten Benutzer"""),

    ("Anforderung_Kalenderverwaltung", "Fachliche Anforderung: Erweiterte Kalenderverwaltung",
     """Fachliche Anforderung: Erweiterte Kalenderverwaltung

1. Hintergrund
Die Terminplanung von Streams beruecksichtigt aktuell nur Werktage und gesetzliche Feiertage. Kundenspezifische Betriebskalender werden nicht unterstuetzt.

2. Anforderungen
- Import von SAP-Fabrikkalendern (T001W)
- Kundenspezifische Feiertage und Betriebsruhetage
- Schichtkalender mit 3-Schicht-Modell
- Ausnahmetage (Betriebsversammlungen, Inventur)
- Kalender-Vererbung (Konzern -> Standort -> Abteilung)

3. Regelwerk
- "Nur an Werktagen": Mo-Fr, kein Feiertag
- "Letzter Werktag des Monats": Beruecksichtigt Feiertage
- "Erster Arbeitstag nach Monatsende": Verschiebt bei Feiertag
- "Alle 14 Tage freitags": Kalenderwoche-basiert

4. Schnittstelle
Kalender sollen ueber eine REST-API verwaltbar sein und per XML importiert/exportiert werden koennen."""),

    ("Anforderung_Mandantenfaehigkeit", "Fachliche Anforderung: Mandantenfaehigkeit",
     """Fachliche Anforderung: Multi-Mandanten-Faehigkeit

1. Geschaeftsanforderung
Streamworks soll fuer mehrere Kunden (Mandanten) auf einer gemeinsamen Infrastruktur betrieben werden koennen.

2. Isolation
- Strikte Datentrennung zwischen Mandanten
- Eigene Konfigurationen pro Mandant
- Separate Berechtigungsstrukturen
- Getrennte Logging- und Audit-Trails

3. Ressourcenmanagement
- Zuweisung von Agent-Kapazitaeten pro Mandant
- Fair-Scheduling bei gemeinsam genutzten Ressourcen
- Priorisierung konfigurierbar pro Mandant
- Ressourcen-Quotas verhindern Uebernutzung

4. Administration
- Mandanten-Admin kann eigene Benutzer verwalten
- Plattform-Admin sieht alle Mandanten
- Self-Service-Onboarding fuer neue Mandanten
- Mandantenspezifische Dashboards und Reports"""),

    ("Anforderung_API_Gateway", "Fachliche Anforderung: REST API Gateway",
     """Fachliche Anforderung: REST API Gateway fuer Streamworks

1. Ziel
Bereitstellung einer REST API fuer die programmatische Steuerung von Streamworks durch Drittsysteme.

2. Funktionsumfang
2.1 Stream-Management
- GET /api/streams - Alle Streams auflisten
- POST /api/streams - Neuen Stream erstellen
- PUT /api/streams/{id} - Stream aktualisieren
- DELETE /api/streams/{id} - Stream loeschen
- POST /api/streams/{id}/execute - Stream manuell starten

2.2 Job-Management
- GET /api/jobs - Jobs auflisten mit Filterung
- GET /api/jobs/{id}/status - Aktuellen Status abfragen
- POST /api/jobs/{id}/cancel - Laufenden Job abbrechen
- GET /api/jobs/{id}/log - Job-Log abrufen

2.3 Monitoring
- GET /api/health - Systemstatus
- GET /api/metrics - Prometheus-kompatible Metriken
- GET /api/agents - Agenten-Status

3. Authentifizierung
- API-Key-basiert fuer Service-to-Service
- OAuth2 fuer Benutzer-Interaktion
- Rate-Limiting: 1000 Requests/Minute pro API-Key

4. Dokumentation
- OpenAPI 3.0 Spezifikation
- Swagger-UI unter /docs
- Beispiel-Requests fuer alle Endpunkte"""),

    ("Anforderung_Archivierung", "Fachliche Anforderung: Log-Archivierung",
     """Fachliche Anforderung: Automatische Log-Archivierung

1. Problem
Job-Logs belegen zunehmend Speicherplatz. Eine manuelle Bereinigung ist aufwaendig und fehleranfaellig.

2. Anforderung
Automatische Archivierung und Bereinigung von Job-Logs nach konfigurierbaren Regeln:
- Online-Aufbewahrung: 30 Tage (konfigurierbar)
- Archiv-Aufbewahrung: 1 Jahr
- Langzeitarchiv: 10 Jahre (nur fuer regulatorisch relevante Jobs)

3. Archivierungsformat
- Komprimierung mit GZIP
- Struktur: /archiv/{jahr}/{monat}/{stream_name}/{job_id}.log.gz
- Metadaten in separater Index-Datei (JSON)

4. Zugriff
- Direkte Suche in Online-Logs (< 30 Tage)
- Archiv-Suche ueber Index (30 Tage - 1 Jahr)
- Langzeitarchiv nur ueber explizite Wiederherstellung

5. Speicherberechnung
- Aktuelle Log-Produktion: ~5 GB/Tag
- Nach Komprimierung: ~500 MB/Tag
- Geschaetzter Jahrsbedarf: ~200 GB"""),

    ("Anforderung_Disaster_Recovery", "Fachliche Anforderung: Disaster Recovery",
     """Fachliche Anforderung: Disaster Recovery und Business Continuity

1. Schutzziele
- RPO (Recovery Point Objective): 1 Stunde
- RTO (Recovery Time Objective): 4 Stunden
- MTPD (Maximum Tolerable Period of Disruption): 8 Stunden

2. Backup-Strategie
- Vollbackup: Woechentlich (Sonntag 02:00 Uhr)
- Inkrementelles Backup: Taeglich um 23:00 Uhr
- Transaktionslog-Sicherung: Stuendlich

3. Failover
- Aktiv-Passiv-Cluster fuer Streamworks-Server
- Automatischer Failover bei Heartbeat-Verlust (> 60 Sekunden)
- Datenbank-Replikation auf Standby-System
- DNS-Failover fuer Client-Zugriff

4. Recovery-Prozedur
Schritt 1: Schadensbewertung (max. 30 Minuten)
Schritt 2: Entscheidung Failover/Reparatur
Schritt 3: Failover aktivieren (automatisch oder manuell)
Schritt 4: Datenkonsistenz pruefen
Schritt 5: Ausstehende Jobs neu einplanen
Schritt 6: Normalbet rieb aufnehmen

5. Tests
- DR-Test quartalsweise
- Failover-Test monatlich
- Backup-Restore-Test woechentlich (automatisiert)"""),

    ("Anforderung_Compliance", "Fachliche Anforderung: Compliance und Audit",
     """Fachliche Anforderung: Compliance und Audit-Trail

1. Regulatorische Anforderungen
Streamworks muss die Anforderungen von SOX (Sarbanes-Oxley), GDPR und branchenspezifischen Regulierungen erfuellen.

2. Audit-Trail
Vollstaendige Protokollierung aller:
- Konfigurationsaenderungen (wer, wann, was, vorher/nachher)
- Job-Ausfuehrungen (Start, Ende, Status, Parameter)
- Benutzeraktionen (Login, Logout, Berechtigungsaenderungen)
- Systemereignisse (Starts, Stops, Fehler)

3. Berichtswesen
- SOX-Compliance-Report (quartalsweise)
- Zugriffsprotokoll (auf Anforderung)
- Aenderungshistorie pro Stream (auf Anforderung)
- Datenverarbeitungsverzeichnis (GDPR Art. 30)

4. Datenschutz
- Personenbezogene Daten in Logs maskieren
- Loeschkonzept fuer personenbezogene Daten
- Auskunftsfaehigkeit ueber verarbeitete Daten
- Privacy by Design in neuen Streams

5. Zertifizierung
Vorbereitung auf ISO 27001 Zertifizierung:
- ISMS-Dokumentation
- Risikobewertung
- Massnahmenkatalog"""),
]

BETRIEB_DOCS = [
    ("Betrieb_Backup_Prozedur", "Backup-Prozedur fuer Streamworks",
     """Betriebsanleitung: Backup-Prozedur fuer Streamworks

Gueltig ab: 2025-01-01
Verantwortlich: IT-Betrieb, Team Infrastruktur

1. Taeglich (automatisiert)
- 23:00 Uhr: Datenbank-Backup (PostgreSQL pg_dump)
  Befehl: pg_dump -Fc streamworks_db > /backup/daily/sw_$(date +%Y%m%d).dump
- 23:30 Uhr: Konfigurationsdateien sichern
  Verzeichnisse: /etc/streamworks/, /opt/streamworks/config/
- 00:00 Uhr: Verifizierung des Backups (Checksumme + Restore-Test)

2. Woechentlich (Sonntag)
- 02:00 Uhr: Vollbackup der gesamten Streamworks-Installation
- Inklusive: Datenbank, Konfiguration, Templates, Logs
- Aufbewahrung: 4 Wochen
- Speicherort: NFS-Share /backup/weekly/

3. Monatlich
- Letzter Sonntag: Backup auf Band/Langzeitspeicher
- Aufbewahrung: 12 Monate
- Offsite-Kopie an Standort 2

4. Restore-Verfahren
Schritt 1: Streamworks-Dienste stoppen
  systemctl stop streamworks
Schritt 2: Datenbank wiederherstellen
  pg_restore -d streamworks_db /backup/daily/sw_YYYYMMDD.dump
Schritt 3: Konfiguration wiederherstellen
  tar xzf /backup/daily/config_YYYYMMDD.tar.gz -C /
Schritt 4: Dienste starten und pruefen
  systemctl start streamworks
  curl http://localhost:8000/health

5. Notfallkontakte
- Bereitschaft IT-Betrieb: +49 123 456 7890
- Eskalation: IT-Leitung +49 123 456 7891"""),

    ("Betrieb_Incident_Response", "Incident Response Verfahren",
     """Betriebsanleitung: Incident Response fuer Streamworks

1. Incident-Klassifizierung
Prio 1 (Kritisch): Produktionsstopp, Datenverlust
- Reaktionszeit: 15 Minuten
- Loesungszeit: 4 Stunden
- Eskalation: Sofort an IT-Leitung

Prio 2 (Hoch): Einzelner kritischer Stream ausgefallen
- Reaktionszeit: 30 Minuten
- Loesungszeit: 8 Stunden
- Eskalation: Nach 2 Stunden

Prio 3 (Mittel): Performance-Degradation, Warnungen
- Reaktionszeit: 2 Stunden
- Loesungszeit: 24 Stunden
- Eskalation: Nach 8 Stunden

Prio 4 (Niedrig): Kosmetische Fehler, Feature-Requests
- Reaktionszeit: 1 Arbeitstag
- Loesungszeit: 5 Arbeitstage

2. Erstmassnahmen bei Prio 1
a) Situation erfassen: Was ist ausgefallen? Seit wann?
b) Auswirkung bewerten: Welche Geschaeftsprozesse betroffen?
c) Sofortmassnahme: Workaround moeglich? Manuelle Verarbeitung?
d) Kommunikation: Stakeholder informieren (Verteiler: SW-INCIDENT)
e) Dokumentation: Incident-Ticket im ServiceDesk erfassen

3. Haeufige Incidents und Loesungen
- Agent nicht erreichbar: Neustart des Agent-Dienstes
- Datenbank-Verbindung verloren: Connection-Pool pruefen
- Disk voll: Alte Logs bereinigen, Archivierung anstosse
- SAP-RFC-Timeout: SAP-Gateway pruefen, RFC-Verbindung testen"""),

    ("Betrieb_Agent_Installation", "Agent-Installationsanleitung",
     """Betriebsanleitung: Streamworks Agent Installation

1. Voraussetzungen
- Betriebssystem: Windows Server 2019+ oder RHEL 8+/SLES 15+
- RAM: mindestens 2 GB frei
- Disk: mindestens 10 GB frei
- Netzwerk: Port 8443 (HTTPS) zum Streamworks-Server

2. Windows-Installation
Schritt 1: Installer herunterladen
  https://intranet.firma.de/streamworks/agent/sw-agent-latest.msi

Schritt 2: Installation ausfuehren
  msiexec /i sw-agent-latest.msi SERVERURL=https://streamworks.firma.de

Schritt 3: Dienst pruefen
  sc query StreamworksAgent

Schritt 4: Agent registrieren
  Im Streamworks-Portal unter Administration > Agenten > Registrierung

3. Linux-Installation
Schritt 1: Repository einrichten
  echo "deb https://repo.firma.de/streamworks stable main" > /etc/apt/sources.list.d/streamworks.list

Schritt 2: Installieren
  apt-get update && apt-get install streamworks-agent

Schritt 3: Konfigurieren
  vi /etc/streamworks/agent.conf
  -> server_url=https://streamworks.firma.de
  -> agent_name=$(hostname)

Schritt 4: Dienst starten
  systemctl enable --now streamworks-agent

4. Fehlerbehebung
- Agent verbindet nicht: Firewall-Regeln pruefen (Port 8443)
- Zertifikatsfehler: CA-Zertifikat importieren
- Performance-Probleme: Java-Heap-Groesse anpassen (-Xmx)"""),

    ("Betrieb_Wartungsfenster", "Wartungsfenster-Prozedur",
     """Betriebsanleitung: Wartungsfenster-Durchfuehrung

1. Regelmaessige Wartungsfenster
- Woechentlich: Sonntag 02:00-06:00 Uhr (automatisch)
- Monatlich: Letzter Sonntag 00:00-08:00 Uhr (geplant)
- Quartalsweise: Upgrade-Fenster nach Abstimmung

2. Vorbereitung (T-1 Tag)
Checkliste:
[ ] Wartungsankuendigung an Stakeholder versendet
[ ] Backup der aktuellen Konfiguration erstellt
[ ] Rollback-Plan dokumentiert und getestet
[ ] Alle laufenden Streams abgeschlossen
[ ] Neue Version auf Test-System validiert

3. Durchfuehrung
Schritt 1: Wartungsmodus aktivieren
  sw-admin maintenance enable --message "Geplante Wartung bis 06:00 Uhr"

Schritt 2: Laufende Jobs abwarten
  sw-admin jobs wait --timeout 30m

Schritt 3: Update einspielen
  sw-admin upgrade --version X.Y.Z

Schritt 4: Konfiguration anpassen (falls noetig)
  sw-admin config apply /path/to/new-config.yaml

Schritt 5: Systemtest durchfuehren
  sw-admin test run --suite post-upgrade

Schritt 6: Wartungsmodus beenden
  sw-admin maintenance disable

4. Rollback bei Problemen
  sw-admin upgrade rollback --to-version X.Y.Z-1
  sw-admin test run --suite post-rollback"""),

    ("Betrieb_Performance_Tuning", "Performance-Tuning Richtlinien",
     """Betriebsanleitung: Performance-Tuning fuer Streamworks

1. Datenbank-Optimierung
1.1 PostgreSQL-Einstellungen
- shared_buffers: 25% des RAM (z.B. 4GB bei 16GB RAM)
- effective_cache_size: 75% des RAM
- work_mem: 256MB fuer komplexe Queries
- maintenance_work_mem: 1GB fuer VACUUM/INDEX

1.2 Indexierung
Wichtige Indizes fuer Streamworks:
- idx_jobs_status ON jobs(status) WHERE status = 'running'
- idx_jobs_stream_date ON jobs(stream_id, scheduled_date)
- idx_logs_job_created ON job_logs(job_id, created_at)

2. Agenten-Optimierung
- Maximale parallele Jobs pro Agent: CPU-Kerne * 2
- Java-Heap: 50% des verfuegbaren RAM
- GC-Algorithmus: G1GC fuer Agenten mit > 4GB Heap
- Thread-Pool-Groesse: 50 (Standard), 100 (Hochlast)

3. Netzwerk
- Keep-Alive-Timeout: 300 Sekunden
- Connection-Pool-Groesse: 20 pro Agent
- Komprimierung: GZIP fuer Payloads > 1KB
- WebSocket-Heartbeat: 30 Sekunden

4. Monitoring
Kritische Metriken ueberwachen:
- CPU-Auslastung Server/Agent > 80%
- Speicherauslastung > 85%
- Disk I/O-Wait > 20%
- Datenbankverbindungen > 80% des Pools
- Job-Queue-Laenge > 100"""),

    ("Betrieb_Loganalyse", "Log-Analyse und Troubleshooting",
     """Betriebsanleitung: Log-Analyse und Troubleshooting

1. Log-Dateien
Streamworks erzeugt folgende Log-Dateien:
- /var/log/streamworks/server.log - Hauptlog des Servers
- /var/log/streamworks/agent-{name}.log - Agent-Logs
- /var/log/streamworks/job-{id}.log - Individuelle Job-Logs
- /var/log/streamworks/audit.log - Audit-Trail

2. Log-Level
- ERROR: Kritische Fehler, erfordern Aktion
- WARN: Potentielle Probleme, Ueberwachung empfohlen
- INFO: Normale Betriebsmeldungen
- DEBUG: Detaillierte Informationen (nur bei Bedarf aktivieren)

3. Haeufige Fehlermuster
3.1 "Connection refused to agent"
Ursache: Agent-Dienst nicht gestartet oder Firewall blockiert
Loesung: Agent-Status pruefen, Netzwerkverbindung testen

3.2 "Job execution timeout after 3600s"
Ursache: Job laeuft laenger als erlaubt
Loesung: Timeout erhoehen oder Job-Logik optimieren

3.3 "Database connection pool exhausted"
Ursache: Zu viele gleichzeitige Datenbankzugriffe
Loesung: Pool-Groesse erhoehen, Slow Queries optimieren

4. Log-Rotation
- Taeglich rotieren bei > 100MB
- Komprimierung nach 1 Tag
- Aufbewahrung: 30 Tage online, 1 Jahr Archiv"""),

    ("Betrieb_Sicherheitsrichtlinien", "Sicherheitsrichtlinien",
     """Betriebsanleitung: Sicherheitsrichtlinien fuer Streamworks

1. Zugangssteuerung
- Alle Zugriffe ueber HTTPS (TLS 1.2+)
- Authentifizierung ueber LDAP/Active Directory
- Zwei-Faktor-Authentifizierung fuer Admin-Zugang
- Session-Timeout: 30 Minuten Inaktivitaet

2. Berechtigungskonzept
Rollen:
- Viewer: Nur-Lese-Zugriff auf Streams und Logs
- Operator: Starten/Stoppen von Streams, Log-Zugriff
- Developer: Stream-Erstellung und -Aenderung
- Admin: Vollzugriff inkl. Benutzerverwaltung

3. Netzwerksicherheit
- Streamworks-Server: Nur Port 8443 (HTTPS) offen
- Agenten: Nur ausgehende Verbindung zum Server
- Datenbank: Nur vom Server erreichbar
- Management-Netzwerk: Separates VLAN

4. Passwort-Richtlinie
- Mindestlaenge: 12 Zeichen
- Komplexitaet: Gross/Klein, Zahl, Sonderzeichen
- Aenderung: Alle 90 Tage
- Historie: Letzte 12 Passwoerter gesperrt

5. Vulnerability Management
- Monatliche Sicherheitsupdates
- Quartalsweise Penetrationstests
- Jaehrlicher Security-Audit"""),

    ("Betrieb_Kapazitaetsplanung", "Kapazitaetsplanung",
     """Betriebsanleitung: Kapazitaetsplanung fuer Streamworks

1. Aktuelle Kapazitaet
- Server: 2x (Aktiv-Passiv), 16 Kerne, 64 GB RAM
- Agenten: 8x Windows, 4x Linux
- Datenbank: PostgreSQL 15, 500 GB SSD
- Parallele Jobs: Max. 200

2. Wachstumsprognose
Aktuelle Nutzung und Trend:
- Jobs/Tag: 5.000 (Wachstum: +15%/Jahr)
- Streams: 350 (Wachstum: +25 neue/Quartal)
- Agenten-Auslastung: 65% (Peak: 85%)
- Datenbankgroesse: 180 GB (Wachstum: +5 GB/Monat)

3. Schwellwerte fuer Erweiterung
- CPU-Auslastung dauerhaft > 70%: Neuen Agent hinzufuegen
- RAM-Auslastung > 80%: Server-RAM erhoehen
- Datenbank > 80% der Disk: Speicher erweitern
- Job-Queue > 50 wartende Jobs: Parallelitaet erhoehen

4. Planungshorizont
- 6 Monate: Neuer Agent fuer geplante SAP S/4HANA Migration
- 12 Monate: Datenbank-Upgrade auf 1 TB
- 18 Monate: Evaluierung Cluster-Architektur
- 24 Monate: Cloud-Migration evaluieren"""),

    ("Betrieb_Notfallhandbuch", "Notfallhandbuch",
     """Betriebsanleitung: Notfallhandbuch Streamworks

WICHTIG: Dieses Handbuch ist auch offline verfuegbar unter:
/opt/streamworks/docs/notfall/ und als gedruckte Version im Serverraum.

1. Totalausfall Streamworks-Server
Sofortmassnahme:
a) Failover pruefen: Ist Standby-Server aktiv?
   sw-admin cluster status
b) Falls nicht: Manueller Failover
   ssh standby "sw-admin cluster promote"
c) DNS pruefen: Zeigt auf aktiven Server?

2. Totalausfall Datenbank
a) Replikationsstatus pruefen
   psql -c "SELECT * FROM pg_stat_replication;"
b) Failover auf Standby-DB
   pg_ctl promote -D /var/lib/postgresql/data
c) Streamworks-Config anpassen
   sw-admin config set db.host=standby-db.firma.de

3. Netzwerkausfall
a) Agenten arbeiten im Offline-Modus weiter
b) Jobs werden lokal gepuffert
c) Nach Wiederherstellung: Automatische Synchronisation

4. Stromausfall
a) USV versorgt Server fuer 30 Minuten
b) Nach 15 Minuten: Automatischer kontrollierter Shutdown
c) Nach Stromrueckkehr: Automatischer Start via IPMI/iLO
d) Streamworks startet automatisch und nimmt ausstehende Jobs wieder auf

5. Wichtige Telefonnummern
- IT-Bereitschaft: +49 123 456 7890
- Rechenzentrumsbetrieb: +49 123 456 7895
- SAP-Basis-Team: +49 123 456 7892
- Netzwerk-Team: +49 123 456 7893
- IT-Sicherheit: +49 123 456 7894"""),

    ("Betrieb_Update_Checkliste", "Update-Checkliste",
     """Betriebsanleitung: Update-Checkliste fuer Streamworks

Version: Allgemein gueltig
Letzte Aktualisierung: 2025-10-01

CHECKLISTE VOR DEM UPDATE
=========================

[ ] 1. Release Notes gelesen und verstanden
[ ] 2. Breaking Changes identifiziert
[ ] 3. Kompatibilitaet mit Agenten-Version geprueft
[ ] 4. Update auf Test-System erfolgreich durchgefuehrt
[ ] 5. Regressionstests auf Test-System bestanden
[ ] 6. Backup des Produktivsystems erstellt
[ ] 7. Rollback-Plan dokumentiert
[ ] 8. Wartungsfenster genehmigt und kommuniziert
[ ] 9. Bereitschaftsplan fuer Update-Zeitraum erstellt
[ ] 10. Monitoring-Schwellwerte angepasst

CHECKLISTE WAEHREND DES UPDATES
================================

[ ] 1. Wartungsmodus aktiviert
[ ] 2. Laufende Jobs abgewartet
[ ] 3. Datenbank-Migration ausgefuehrt
[ ] 4. Server-Software aktualisiert
[ ] 5. Konfiguration migriert
[ ] 6. Agenten aktualisiert (rolling update)
[ ] 7. Smoke-Tests bestanden
[ ] 8. Wartungsmodus deaktiviert

CHECKLISTE NACH DEM UPDATE
===========================

[ ] 1. Alle Streams gestartet und aktiv
[ ] 2. Agenten verbunden und gesund
[ ] 3. Erste Job-Ausfuehrungen erfolgreich
[ ] 4. Performance-Baseline erstellt
[ ] 5. Stakeholder ueber erfolgreichen Abschluss informiert
[ ] 6. Update-Dokumentation archiviert"""),
]

CONFIG_DOCS = [
    ("Config_Agent_Mapping", "Agent-Zuordnungstabelle",
     [["Agent-Name", "Hostname", "Betriebssystem", "SAP-System", "Max Jobs", "Standort", "Status"],
      ["WIN-AGENT-001", "srvwin001.firma.de", "Windows 2019", "PRD", "20", "Hamburg", "Aktiv"],
      ["WIN-AGENT-002", "srvwin002.firma.de", "Windows 2019", "PRD", "20", "Hamburg", "Aktiv"],
      ["WIN-AGENT-003", "srvwin003.firma.de", "Windows 2022", "QAS", "15", "Hamburg", "Aktiv"],
      ["WIN-AGENT-004", "srvwin004.firma.de", "Windows 2022", "DEV", "10", "Muenchen", "Aktiv"],
      ["LNX-AGENT-001", "srvlnx001.firma.de", "RHEL 8", "PRD", "25", "Hamburg", "Aktiv"],
      ["LNX-AGENT-002", "srvlnx002.firma.de", "RHEL 8", "PRD", "25", "Hamburg", "Aktiv"],
      ["LNX-AGENT-003", "srvlnx003.firma.de", "SLES 15", "QAS", "20", "Berlin", "Aktiv"],
      ["LNX-AGENT-004", "srvlnx004.firma.de", "SLES 15", "DEV", "15", "Berlin", "Wartung"],
      ["UNX-AGENT-001", "srvunx001.firma.de", "AIX 7.2", "PRD", "10", "Hamburg", "Aktiv"],
      ["UNX-AGENT-002", "srvunx002.firma.de", "AIX 7.2", "PRD", "10", "Hamburg", "Aktiv"]]),

    ("Config_Schedule_Matrix", "Scheduling-Matrix",
     [["Stream-Kategorie", "Frequenz", "Startzeit", "Prioritaet", "Max Laufzeit", "Recovery", "SLA"],
      ["SAP-Extraktion", "Taeglich", "23:00", "Hoch", "4h", "3x Retry", "99.5%"],
      ["SAP-Buchungen", "Taeglich", "06:00", "Kritisch", "2h", "Sofort-Alarm", "99.9%"],
      ["Dateitransfer", "Alle 2h", "00:00", "Mittel", "1h", "3x Retry", "98%"],
      ["Monitoring", "Alle 15min", "Laufend", "Niedrig", "5min", "Neustart", "95%"],
      ["Backup", "Taeglich", "02:00", "Hoch", "3h", "Manuell", "99%"],
      ["Reporting", "Taeglich", "07:00", "Mittel", "2h", "Skip", "95%"],
      ["Archivierung", "Woechentlich", "Sonntag 03:00", "Niedrig", "6h", "Naechste Woche", "90%"],
      ["Monatsabschluss", "Monatlich", "Letzter Werktag", "Kritisch", "8h", "Manuell", "99.9%"],
      ["Batch-Input", "Taeglich", "21:00", "Hoch", "3h", "3x Retry", "99%"],
      ["IDoc-Verarbeitung", "Alle 30min", "Laufend", "Hoch", "15min", "Neustart", "99.5%"]]),

    ("Config_SAP_Systeme", "SAP-Systemlandschaft",
     [["SID", "Typ", "Hostname", "Instanz", "Client", "Mandant", "RFC-Destination", "Status"],
      ["PRD", "Produktion", "sapprd01", "00", "100", "Produktiv", "PRD_RFC_001", "Aktiv"],
      ["QAS", "Qualitaet", "sapqas01", "00", "200", "Test", "QAS_RFC_001", "Aktiv"],
      ["DEV", "Entwicklung", "sapdev01", "00", "300", "Entwicklung", "DEV_RFC_001", "Aktiv"],
      ["SBX", "Sandbox", "sapsbx01", "00", "400", "Spielwiese", "SBX_RFC_001", "Inaktiv"],
      ["BW1", "BW Produktion", "sapbw01", "00", "100", "BW Prod", "BW1_RFC_001", "Aktiv"],
      ["BW2", "BW Qualitaet", "sapbw02", "00", "200", "BW Test", "BW2_RFC_001", "Aktiv"],
      ["PI1", "PI Produktion", "sappi01", "00", "100", "PI Prod", "PI1_RFC_001", "Aktiv"],
      ["SM1", "SolMan", "sapsolman", "00", "100", "SolMan", "SM1_RFC_001", "Aktiv"],
      ["GRC", "GRC System", "sapgrc01", "00", "100", "Governance", "GRC_RFC_001", "Aktiv"],
      ["S4H", "S/4HANA Pilot", "saps4h01", "00", "100", "S4 Pilot", "S4H_RFC_001", "Planung"]]),

    ("Config_Berechtigungsmatrix", "Berechtigungsmatrix",
     [["Rolle", "Streams lesen", "Streams erstellen", "Streams aendern", "Streams loeschen", "Jobs starten", "Admin", "Audit"],
      ["Viewer", "Ja", "Nein", "Nein", "Nein", "Nein", "Nein", "Nein"],
      ["Operator", "Ja", "Nein", "Nein", "Nein", "Ja", "Nein", "Nein"],
      ["Developer", "Ja", "Ja", "Ja", "Nein", "Ja", "Nein", "Nein"],
      ["Team-Lead", "Ja", "Ja", "Ja", "Ja", "Ja", "Nein", "Ja"],
      ["Admin", "Ja", "Ja", "Ja", "Ja", "Ja", "Ja", "Ja"],
      ["Auditor", "Ja", "Nein", "Nein", "Nein", "Nein", "Nein", "Ja"],
      ["Service-Account", "Ja", "Nein", "Nein", "Nein", "Ja", "Nein", "Nein"]]),

    ("Config_Fehlercode_Katalog", "Fehlercode-Katalog",
     [["Code", "Kategorie", "Beschreibung", "Schwere", "Aktion", "Auto-Recovery"],
      ["SW-001", "Verbindung", "Agent nicht erreichbar", "Kritisch", "Agent-Neustart pruefen", "Ja"],
      ["SW-002", "Verbindung", "Datenbank-Timeout", "Hoch", "Connection-Pool pruefen", "Ja"],
      ["SW-003", "Verbindung", "SAP RFC fehlgeschlagen", "Hoch", "RFC-Destination pruefen", "Ja (3x)"],
      ["SW-010", "Ausfuehrung", "Job-Timeout ueberschritten", "Mittel", "Timeout erhoehen", "Nein"],
      ["SW-011", "Ausfuehrung", "Job abgebrochen (Benutzer)", "Niedrig", "Keine", "Nein"],
      ["SW-012", "Ausfuehrung", "Rueckgabewert != 0", "Mittel", "Job-Log pruefen", "Nein"],
      ["SW-020", "Ressource", "Disk-Space < 10%", "Hoch", "Bereinigung starten", "Ja"],
      ["SW-021", "Ressource", "RAM-Auslastung > 90%", "Hoch", "Prozesse pruefen", "Nein"],
      ["SW-030", "Berechtigung", "Authentifizierung fehlgeschlagen", "Mittel", "Credentials pruefen", "Nein"],
      ["SW-031", "Berechtigung", "Keine Ausfuehrungsberechtigung", "Niedrig", "Rollenzuweisung pruefen", "Nein"],
      ["SW-040", "Datei", "Datei nicht gefunden", "Mittel", "Pfad pruefen", "Nein"],
      ["SW-041", "Datei", "Datei gesperrt", "Mittel", "Sperrenden Prozess pruefen", "Ja (5min Warten)"]]),

    ("Config_SLA_Definitionen", "SLA-Definitionen",
     [["SLA-Klasse", "Verfuegbarkeit", "Max Ausfallzeit/Monat", "Reaktionszeit", "Loesungszeit", "Eskalation nach"],
      ["Platin", "99.99%", "4.3 Minuten", "5 Minuten", "1 Stunde", "15 Minuten"],
      ["Gold", "99.9%", "43.2 Minuten", "15 Minuten", "4 Stunden", "1 Stunde"],
      ["Silber", "99.5%", "3.6 Stunden", "30 Minuten", "8 Stunden", "2 Stunden"],
      ["Bronze", "99%", "7.2 Stunden", "2 Stunden", "24 Stunden", "8 Stunden"],
      ["Standard", "95%", "36 Stunden", "4 Stunden", "5 Arbeitstage", "24 Stunden"]]),

    ("Config_Netzwerk_Ports", "Netzwerk-Port-Matrix",
     [["Dienst", "Port", "Protokoll", "Richtung", "Quelle", "Ziel", "Firewall-Regel"],
      ["Streamworks UI", "8443", "HTTPS", "Eingehend", "Client-Netz", "SW-Server", "FW-SW-001"],
      ["Streamworks API", "8443", "HTTPS", "Eingehend", "API-Consumer", "SW-Server", "FW-SW-002"],
      ["Agent-Kommunikation", "8443", "HTTPS", "Ausgehend", "Agenten", "SW-Server", "FW-SW-003"],
      ["PostgreSQL", "5432", "TCP", "Intern", "SW-Server", "DB-Server", "FW-SW-010"],
      ["Qdrant", "6333", "HTTP", "Intern", "SW-Server", "Qdrant-Server", "FW-SW-011"],
      ["MinIO", "9000", "HTTP", "Intern", "SW-Server", "MinIO-Server", "FW-SW-012"],
      ["MinIO Console", "9001", "HTTPS", "Eingehend", "Admin-Netz", "MinIO-Server", "FW-SW-013"],
      ["SAP RFC", "3300", "TCP", "Ausgehend", "Agenten", "SAP-System", "FW-SAP-001"],
      ["SMTP", "587", "TLS", "Ausgehend", "SW-Server", "Mail-Relay", "FW-SW-020"],
      ["LDAP", "636", "LDAPS", "Ausgehend", "SW-Server", "AD-Server", "FW-SW-021"]]),

    ("Config_Umgebungsvariablen", "Umgebungsvariablen-Referenz",
     [["Variable", "Pflicht", "Default", "Beschreibung", "Beispiel"],
      ["OPENAI_API_KEY", "Ja", "-", "OpenAI API Schluessel", "sk-..."],
      ["SUPABASE_URL", "Nein", "In-Memory", "Supabase Projekt-URL", "https://xxx.supabase.co"],
      ["SUPABASE_KEY", "Nein", "In-Memory", "Supabase Anon Key", "eyJ..."],
      ["QDRANT_URL", "Nein", "http://localhost:6333", "Qdrant Verbindungs-URL", "http://qdrant:6333"],
      ["QDRANT_COLLECTION", "Nein", "streamworks", "Qdrant Collection Name", "streamworks"],
      ["MINIO_ENDPOINT", "Nein", "localhost:9000", "MinIO Server Adresse", "minio:9000"],
      ["MINIO_ACCESS_KEY", "Nein", "streamworks", "MinIO Zugangsdaten", "streamworks"],
      ["MINIO_SECRET_KEY", "Nein", "streamworks123", "MinIO Passwort", "sicheresPasswort"],
      ["MINIO_BUCKET", "Nein", "documents", "MinIO Bucket Name", "documents"],
      ["BACKEND_HOST", "Nein", "0.0.0.0", "Backend Bind-Adresse", "0.0.0.0"],
      ["BACKEND_PORT", "Nein", "8000", "Backend Port", "8000"],
      ["CORS_ORIGINS", "Nein", "localhost:3000", "Erlaubte Origins", "http://localhost:3000"],
      ["ENVIRONMENT", "Nein", "development", "Umgebungstyp", "production"]]),

    ("Config_Job_Typen", "Job-Typ-Konfiguration",
     [["Job-Typ", "Beschreibung", "Agent-OS", "Typischer Timeout", "Recovery-Strategie", "Parameter"],
      ["SAP-Job", "SAP Background Job via XBP", "Windows/Linux", "3600s", "3x Retry + Alarm", "SID, Client, Variante"],
      ["Script-Job", "Shell/PowerShell Script", "Passend zum OS", "1800s", "2x Retry", "Script-Pfad, Argumente"],
      ["File-Transfer", "SFTP/FTP Dateiuebertragung", "Alle", "900s", "3x Retry", "Quelle, Ziel, Muster"],
      ["File-Watch", "Dateiueberwachung", "Alle", "Unbegrenzt", "Neustart", "Pfad, Muster, Aktion"],
      ["Database-Job", "SQL-Ausfuehrung", "Alle", "3600s", "Kein Retry", "Connection, SQL, Parameter"],
      ["Mail-Job", "E-Mail versenden", "Alle", "120s", "2x Retry", "Empfaenger, Betreff, Template"],
      ["REST-Call", "HTTP API Aufruf", "Alle", "300s", "3x Retry", "URL, Methode, Payload"],
      ["Batch-Input", "SAP Batch-Input Verarbeitung", "Windows", "7200s", "Manuell", "Mappe, Modus"]]),
]

BERICHT_DOCS = [
    ("Bericht_Jobausfuehrung_Q4", "Job-Ausfuehrungsstatistik Q4 2025", "csv",
     [["Monat", "Jobs gesamt", "Erfolgreich", "Fehlgeschlagen", "Abgebrochen", "Erfolgsquote", "Durchschn. Laufzeit"],
      ["Oktober 2025", "152340", "150120", "1890", "330", "98.54%", "2m 34s"],
      ["November 2025", "148900", "147200", "1420", "280", "98.86%", "2m 28s"],
      ["Dezember 2025", "138500", "136800", "1450", "250", "98.77%", "2m 45s"],
      ["Gesamt Q4", "439740", "434120", "4760", "860", "98.72%", "2m 36s"]]),

    ("Bericht_Performance_Agents", "Agent-Performance Report", "csv",
     [["Agent", "Jobs ausgefuehrt", "CPU Durchschn.", "RAM Durchschn.", "Disk I/O", "Fehlerquote", "Uptime"],
      ["WIN-AGENT-001", "18500", "45%", "62%", "35 MB/s", "1.2%", "99.95%"],
      ["WIN-AGENT-002", "17800", "42%", "58%", "32 MB/s", "0.9%", "99.98%"],
      ["WIN-AGENT-003", "12300", "38%", "52%", "28 MB/s", "1.5%", "99.90%"],
      ["LNX-AGENT-001", "22100", "55%", "48%", "45 MB/s", "0.6%", "99.99%"],
      ["LNX-AGENT-002", "21500", "52%", "46%", "42 MB/s", "0.7%", "99.99%"],
      ["LNX-AGENT-003", "15600", "40%", "44%", "35 MB/s", "1.1%", "99.92%"],
      ["UNX-AGENT-001", "8900", "35%", "55%", "25 MB/s", "0.8%", "99.97%"],
      ["UNX-AGENT-002", "8200", "33%", "53%", "23 MB/s", "0.9%", "99.96%"]]),

    ("Bericht_SLA_Einhaltung", "SLA-Compliance Report", "csv",
     [["Stream-Kategorie", "SLA-Klasse", "SLA-Ziel", "Ist-Wert", "Eingehalten", "Verletzte SLAs", "Ursache"],
      ["SAP-Buchungen", "Platin", "99.99%", "99.98%", "Nein", "1", "Netzwerk-Unterbrechung 03.12."],
      ["SAP-Extraktion", "Gold", "99.9%", "99.95%", "Ja", "0", "-"],
      ["Dateitransfer", "Silber", "99.5%", "99.7%", "Ja", "0", "-"],
      ["Monitoring", "Bronze", "99%", "99.8%", "Ja", "0", "-"],
      ["Reporting", "Bronze", "99%", "99.5%", "Ja", "0", "-"],
      ["Monatsabschluss", "Platin", "99.99%", "100%", "Ja", "0", "-"]]),

    ("Bericht_Kapazitaetsauslastung", "Kapazitaetsauslastung Dezember 2025", "pdf",
     """Kapazitaetsauslastungsbericht - Dezember 2025

1. Zusammenfassung
Die Gesamtauslastung der Streamworks-Infrastruktur lag im Dezember 2025 bei durchschnittlich 68%. Spitzenwerte wurden am Monatsende (Jahresabschluss-Verarbeitung) mit 92% erreicht.

2. Server-Auslastung
- CPU: Durchschnitt 52%, Peak 85% (31.12. um 23:00)
- RAM: Durchschnitt 61%, Peak 78%
- Disk I/O: Durchschnitt 35%, Peak 72%
- Netzwerk: Durchschnitt 28%, Peak 55%

3. Agenten-Auslastung
Die 12 Agenten waren durchschnittlich zu 65% ausgelastet. Der Bottleneck lag bei den Windows-Agenten (72%), waehrend Linux-Agenten bei 58% lagen.

Empfehlung: Zwei zusaetzliche Windows-Agenten fuer Q1 2026 einplanen, um Reserve fuer die S/4HANA-Migration zu schaffen.

4. Datenbank
- Tabellengroesse: 182 GB (Wachstum: +4.8 GB im Dezember)
- Aktive Verbindungen: Durchschnitt 45, Peak 112
- Langsamste Query: 4.2s (idx_jobs_status Rebuild empfohlen)

5. Prognose Q1 2026
Basierend auf dem aktuellen Wachstumstrend wird die Kapazitaet bis Maerz 2026 ausreichen. Ab April wird ein zusaetzlicher Agenten-Server benoetigt."""),

    ("Bericht_Fehleranalyse_Dezember", "Fehleranalyse Dezember 2025", "pdf",
     """Fehleranalysebericht - Dezember 2025

1. Uebersicht
Im Dezember 2025 traten insgesamt 1.450 Job-Fehler auf (Fehlerquote: 1.05%). Dies liegt unter dem SLA-Ziel von 1.5%.

2. Top-5 Fehlerursachen
1. SAP RFC Timeout (SW-003): 380 Fehler (26.2%)
   - Hauptsaechlich am 15.12. waehrend SAP-Wartung
   - Empfehlung: Wartungsfenster-Erkennung implementieren

2. Datei nicht gefunden (SW-040): 290 Fehler (20.0%)
   - Ursache: Verzoegerte Bereitstellung durch Vorsystem
   - Empfehlung: Wartezeit auf File-Watch erhoehen

3. Job-Timeout (SW-010): 185 Fehler (12.8%)
   - Betroffene Streams: Monatsabschluss, BW-Extraktion
   - Empfehlung: Timeout fuer Monatsende-Jobs erhoehen

4. Agent nicht erreichbar (SW-001): 140 Fehler (9.7%)
   - LNX-AGENT-004 waehrend Wartung offline
   - Empfehlung: Agent-Wartung besser koordinieren

5. Disk-Space-Warnung (SW-020): 95 Fehler (6.6%)
   - WIN-AGENT-003: Temp-Verzeichnis voll
   - Empfehlung: Automatische Temp-Bereinigung einrichten

3. Massnahmen
- RFC-Timeout auf 600s erhoehen (umgesetzt am 20.12.)
- Wartungsfenster-Kalender fuer Agenten einrichten (geplant Q1)
- Automatische Temp-Bereinigung implementieren (geplant Januar)"""),

    ("Bericht_Audit_Q4", "Audit-Trail Zusammenfassung Q4", "csv",
     [["Datum", "Benutzer", "Aktion", "Objekt", "Detail", "IP-Adresse"],
      ["2025-10-01", "admin01", "Login", "System", "Erfolgreiche Anmeldung", "10.0.1.50"],
      ["2025-10-02", "dev03", "Stream erstellt", "GECK003_NEU_001", "Neuer SAP-Job-Stream", "10.0.2.30"],
      ["2025-10-05", "op01", "Job gestartet", "GECK003_FI_MONAT", "Manueller Start", "10.0.1.55"],
      ["2025-10-10", "admin01", "Benutzer angelegt", "dev05", "Rolle: Developer", "10.0.1.50"],
      ["2025-10-15", "dev03", "Stream geaendert", "GECK003_NEU_001", "Timeout angepasst", "10.0.2.30"],
      ["2025-11-01", "admin01", "Rolle geaendert", "op02", "Operator -> Team-Lead", "10.0.1.50"],
      ["2025-11-10", "dev05", "Stream erstellt", "GECK003_BW_DELTA", "BW Delta-Load", "10.0.2.35"],
      ["2025-11-20", "admin01", "Agent registriert", "LNX-AGENT-005", "Neuer Linux-Agent", "10.0.1.50"],
      ["2025-12-01", "op01", "Wartung gestartet", "System", "Geplante Wartung", "10.0.1.55"],
      ["2025-12-15", "admin01", "Backup erstellt", "Datenbank", "Vor SAP-Wartung", "10.0.1.50"],
      ["2025-12-31", "system", "Archivierung", "Logs", "Q4-Logs archiviert", "127.0.0.1"]]),

    ("Bericht_Dateitransfer_Statistik", "Dateitransfer-Statistik", "csv",
     [["Route", "Dateien/Tag", "Volumen/Tag", "Fehlerquote", "Durchschn. Dauer", "Letzter Fehler"],
      ["Hamburg -> Muenchen", "45", "2.3 GB", "0.2%", "3m 12s", "2025-12-20"],
      ["Hamburg -> Berlin", "38", "1.8 GB", "0.1%", "2m 45s", "2025-12-05"],
      ["Muenchen -> Hamburg", "22", "950 MB", "0.3%", "2m 10s", "2025-12-22"],
      ["Berlin -> Hamburg", "18", "720 MB", "0.1%", "1m 55s", "2025-12-01"],
      ["SAP PRD -> BW1", "120", "8.5 GB", "0.5%", "12m 30s", "2025-12-28"],
      ["Extern (SFTP-IN)", "15", "450 MB", "0.8%", "4m 20s", "2025-12-29"],
      ["Extern (SFTP-OUT)", "8", "280 MB", "0.4%", "3m 05s", "2025-12-18"]]),

    ("Bericht_Stream_Inventar", "Stream-Inventar Uebersicht", "csv",
     [["Stream-Name", "Kategorie", "Agent", "Frequenz", "Letzter Lauf", "Status", "Owner"],
      ["GECK003_SAP_MM_BANF", "SAP-Job", "WIN-AGENT-001", "Taeglich 22:00", "2025-12-31", "Aktiv", "team-sap"],
      ["GECK003_SAP_FI_BUCH", "SAP-Job", "WIN-AGENT-001", "Taeglich 06:00", "2025-12-31", "Aktiv", "team-sap"],
      ["GECK003_SAP_SD_LIEF", "SAP-Job", "WIN-AGENT-002", "Taeglich 06:00", "2025-12-31", "Aktiv", "team-sap"],
      ["GECK003_BW_DELTA", "SAP-Job", "LNX-AGENT-001", "Taeglich 23:00", "2025-12-31", "Aktiv", "team-bi"],
      ["GECK003_FT_HH_MUC", "File-Transfer", "LNX-AGENT-002", "Taeglich 22:00", "2025-12-31", "Aktiv", "team-ops"],
      ["GECK003_FT_HH_BER", "File-Transfer", "LNX-AGENT-002", "Taeglich 22:00", "2025-12-31", "Aktiv", "team-ops"],
      ["GECK003_MON_AGENT", "Monitoring", "LNX-AGENT-001", "Alle 15min", "2025-12-31", "Aktiv", "team-ops"],
      ["GECK003_BAK_DB", "Backup", "LNX-AGENT-003", "Taeglich 02:00", "2025-12-31", "Aktiv", "team-infra"],
      ["GECK003_RPT_DAILY", "Reporting", "WIN-AGENT-003", "Taeglich 07:00", "2025-12-31", "Aktiv", "team-ops"],
      ["GECK003_ARC_LOGS", "Archivierung", "LNX-AGENT-003", "Sonntag 03:00", "2025-12-29", "Aktiv", "team-infra"]]),

    ("Bericht_Wartungsprotokoll", "Wartungsprotokoll 2025", "pdf",
     """Wartungsprotokoll - Jahresuebersicht 2025

1. Durchgefuehrte Wartungen
Insgesamt wurden 52 woechentliche und 12 monatliche Wartungsfenster erfolgreich durchgefuehrt.

2. Version-Upgrades
- Januar: Streamworks 8.1.0 -> 8.2.0 (neue Monitoring-Features)
- April: Streamworks 8.2.0 -> 8.3.0 (Performance-Verbesserungen)
- Juli: Streamworks 8.3.0 -> 9.0.0 (Major Release, neue API)
- Oktober: Streamworks 9.0.0 -> 9.1.0 (Bugfixes, neue Agenten-Version)

3. Ungeplante Wartungen
- 15. Maerz: Notfall-Patch fuer Sicherheitsluecke (CVE-2025-1234)
  Dauer: 45 Minuten, keine Datenverluste
- 20. August: Datenbank-Failover wegen Hardware-Defekt
  Dauer: 12 Minuten, automatischer Failover erfolgreich
- 03. November: Agent-Neustart nach Speicherleck
  Dauer: 5 Minuten, keine Job-Unterbrechung

4. Geplante Massnahmen 2026
- Q1: S/4HANA-Agent-Integration
- Q2: Migration auf PostgreSQL 16
- Q3: Cloud-native Architektur evaluieren
- Q4: Streamworks 10.0 (Multi-Mandanten)

5. Kennzahlen
- Gesamte geplante Downtime: 48 Stunden
- Gesamte ungeplante Downtime: 62 Minuten
- Verfuegbarkeit: 99.99%"""),

    ("Bericht_Schulungsnachweis", "Schulungsnachweise Team", "csv",
     [["Mitarbeiter", "Schulung", "Datum", "Dauer", "Zertifikat", "Naechste Auffrischung"],
      ["Mueller, Anna", "Streamworks Administration", "2025-02-15", "3 Tage", "SW-ADM-2025", "2026-02-15"],
      ["Schmidt, Thomas", "Streamworks Entwicklung", "2025-03-10", "5 Tage", "SW-DEV-2025", "2026-03-10"],
      ["Weber, Lisa", "SAP XBP Integration", "2025-04-05", "2 Tage", "SAP-XBP-2025", "2026-04-05"],
      ["Fischer, Mark", "Streamworks Monitoring", "2025-05-20", "1 Tag", "SW-MON-2025", "2026-05-20"],
      ["Becker, Sarah", "IT-Sicherheit Basics", "2025-01-10", "2 Tage", "SEC-BAS-2025", "2026-01-10"],
      ["Hoffmann, Jan", "Streamworks Administration", "2025-06-15", "3 Tage", "SW-ADM-2025", "2026-06-15"],
      ["Koch, Maria", "ITIL Foundation", "2025-07-01", "3 Tage", "ITIL-FND-2025", "2028-07-01"],
      ["Richter, Peter", "Disaster Recovery", "2025-09-10", "1 Tag", "DR-2025", "2026-09-10"]]),
]


# ── Datei-Generatoren ────────────────────────────────────────────────


def generate_pdf(title: str, content: str, filepath: str):
    """PDF mit pymupdf (fitz) erstellen."""
    import fitz

    doc = fitz.open()
    # Titelseite
    page = doc.new_page(width=595, height=842)  # A4
    page.insert_text(
        (50, 80),
        title,
        fontsize=18,
        fontname="helv",
        color=(0, 0.2, 0.4),
    )
    page.insert_text(
        (50, 110),
        "Streamworks-KI Testdokument",
        fontsize=10,
        fontname="helv",
        color=(0.4, 0.4, 0.4),
    )

    # Inhalt ueber mehrere Seiten
    lines = content.strip().split("\n")
    y = 150
    for line in lines:
        if y > 790:
            page = doc.new_page(width=595, height=842)
            y = 50
        fontsize = 12 if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.")) else 10
        bold = line.startswith(("1.", "2.", "3.", "4.", "5.", "6."))
        page.insert_text(
            (50, y),
            line[:90],  # Zeilenlimit
            fontsize=fontsize,
            fontname="hebo" if bold else "helv",
        )
        y += fontsize + 4

    doc.save(filepath)
    doc.close()


def generate_docx(title: str, content: str, filepath: str):
    """DOCX mit python-docx erstellen."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Titel
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x33, 0x66)

    doc.add_paragraph(
        "Streamworks-KI Testdokument",
        style="Subtitle",
    )

    # Inhalt
    for line in content.strip().split("\n"):
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.")):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip().startswith("[ ]"):
            doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(filepath)


def generate_xlsx(title: str, rows: list[list[str]], filepath: str):
    """XLSX mit openpyxl erstellen."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel limit

    # Header-Styling
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for row_idx, row in enumerate(rows, 1):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical="top")

            if row_idx == 1:
                cell.fill = header_fill
                cell.font = header_font
            else:
                cell.font = Font(name="Calibri", size=10)

    # Spaltenbreiten anpassen
    for col_idx in range(1, len(rows[0]) + 1 if rows else 1):
        max_len = max(
            (len(str(rows[r][col_idx - 1])) for r in range(len(rows)) if col_idx - 1 < len(rows[r])),
            default=10,
        )
        ws.column_dimensions[chr(64 + col_idx) if col_idx <= 26 else "A"].width = min(max_len + 4, 40)

    wb.save(filepath)


def generate_txt(title: str, content: str, filepath: str):
    """TXT-Datei erstellen."""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"{'=' * 70}\n")
        f.write(f" {title}\n")
        f.write(f"{'=' * 70}\n\n")
        f.write(content)


def generate_csv(title: str, rows: list[list[str]], filepath: str):
    """CSV-Datei erstellen."""
    with open(filepath, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerows(rows)


# ── API-Aufrufe ──────────────────────────────────────────────────────

client = httpx.Client(timeout=60.0)


def create_folder(name: str, color: str) -> str:
    """Erstellt einen Ordner ueber die API und gibt die folder_id zurueck."""
    resp = client.post(f"{API}/folders", json={"name": name, "color": color})
    resp.raise_for_status()
    return resp.json()["id"]


def upload_file(filepath: str, folder_id: str, mime_type: str) -> dict:
    """Laedt eine Datei via API hoch."""
    filename = os.path.basename(filepath)
    with open(filepath, "rb") as f:
        resp = client.post(
            f"{API}/upload",
            params={"folder_id": folder_id},
            files={"file": (filename, f, mime_type)},
        )
    resp.raise_for_status()
    return resp.json()


# ── Hauptprogramm ────────────────────────────────────────────────────


def main():
    print("=" * 60)
    print(" Streamworks-KI Testdokument-Generator")
    print("=" * 60)

    # Backend pruefen
    try:
        health = client.get(f"{BACKEND_URL}/health")
        health.raise_for_status()
        print(f"\nBackend erreichbar: {BACKEND_URL}")
    except Exception as e:
        print(f"\nFEHLER: Backend nicht erreichbar ({BACKEND_URL})")
        print(f"  {e}")
        print("\nBitte starte das Backend: cd backend && python main.py")
        sys.exit(1)

    # Output-Verzeichnis
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"Temp-Verzeichnis: {OUTPUT_DIR}\n")

    # ── Phase 1: Ordner erstellen ──
    print("Phase 1: Ordner erstellen...")
    folder_ids: dict[str, str] = {}
    for name, color in FOLDER_COLORS.items():
        try:
            fid = create_folder(name, color)
            folder_ids[name] = fid
            print(f"  + {name} ({fid[:8]}...)")
        except httpx.HTTPStatusError as e:
            # Ordner existiert evtl. schon
            print(f"  ! {name}: {e.response.text}")
            # Versuche ID aus bestehenden Ordnern zu holen
            resp = client.get(f"{API}/folders")
            for folder in resp.json():
                if folder["name"] == name:
                    folder_ids[name] = folder["id"]
                    print(f"    -> Existiert bereits ({folder['id'][:8]}...)")
                    break

    print(f"\n{len(folder_ids)} Ordner bereit.\n")

    # ── Phase 2: Dokumente generieren und hochladen ──
    print("Phase 2: Dokumente generieren und hochladen...")
    stats = {"success": 0, "error": 0, "total_chunks": 0}

    # 2.1 SAP-Dokumentation (PDF)
    print("\n--- SAP-Dokumentation (10 PDFs) ---")
    for filename, title, content in SAP_DOCS:
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.pdf")
        generate_pdf(title, content, filepath)
        try:
            result = upload_file(filepath, folder_ids.get("SAP-Dokumentation", ""), "application/pdf")
            stats["success"] += 1
            stats["total_chunks"] += result.get("chunks", 0)
            print(f"  [{stats['success']:2d}/50] {filename}.pdf -> {result.get('chunks', 0)} Chunks")
        except httpx.HTTPStatusError as e:
            stats["error"] += 1
            print(f"  [FEHLER] {filename}.pdf: {e.response.text[:100]}")

    # 2.2 Anforderungen (DOCX)
    print("\n--- Anforderungen (10 DOCX) ---")
    for filename, title, content in ANFORDERUNGEN_DOCS:
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.docx")
        generate_docx(title, content, filepath)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        try:
            result = upload_file(filepath, folder_ids.get("Anforderungen", ""), mime)
            stats["success"] += 1
            stats["total_chunks"] += result.get("chunks", 0)
            print(f"  [{stats['success']:2d}/50] {filename}.docx -> {result.get('chunks', 0)} Chunks")
        except httpx.HTTPStatusError as e:
            stats["error"] += 1
            print(f"  [FEHLER] {filename}.docx: {e.response.text[:100]}")

    # 2.3 Betriebshandbuecher (TXT)
    print("\n--- Betriebshandbuecher (10 TXT) ---")
    for filename, title, content in BETRIEB_DOCS:
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.txt")
        generate_txt(title, content, filepath)
        try:
            result = upload_file(filepath, folder_ids.get("Betriebshandbuecher", ""), "text/plain")
            stats["success"] += 1
            stats["total_chunks"] += result.get("chunks", 0)
            print(f"  [{stats['success']:2d}/50] {filename}.txt -> {result.get('chunks', 0)} Chunks")
        except httpx.HTTPStatusError as e:
            stats["error"] += 1
            print(f"  [FEHLER] {filename}.txt: {e.response.text[:100]}")

    # 2.4 Konfigurationen (XLSX)
    print("\n--- Konfigurationen (10 XLSX) ---")
    for item in CONFIG_DOCS:
        filename, title, rows = item[0], item[1], item[2]
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.xlsx")
        generate_xlsx(title, rows, filepath)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        try:
            result = upload_file(filepath, folder_ids.get("Konfigurationen", ""), mime)
            stats["success"] += 1
            stats["total_chunks"] += result.get("chunks", 0)
            print(f"  [{stats['success']:2d}/50] {filename}.xlsx -> {result.get('chunks', 0)} Chunks")
        except httpx.HTTPStatusError as e:
            stats["error"] += 1
            print(f"  [FEHLER] {filename}.xlsx: {e.response.text[:100]}")

    # 2.5 Berichte (CSV + PDF gemischt)
    print("\n--- Berichte (10 CSV/PDF gemischt) ---")
    for item in BERICHT_DOCS:
        filename, title, fmt = item[0], item[1], item[2]
        data = item[3]

        if fmt == "csv":
            filepath = os.path.join(OUTPUT_DIR, f"{filename}.csv")
            generate_csv(title, data, filepath)
            mime = "text/csv"
        elif fmt == "pdf":
            filepath = os.path.join(OUTPUT_DIR, f"{filename}.pdf")
            generate_pdf(title, data, filepath)
            mime = "application/pdf"
        else:
            continue

        try:
            result = upload_file(filepath, folder_ids.get("Berichte", ""), mime)
            stats["success"] += 1
            stats["total_chunks"] += result.get("chunks", 0)
            print(f"  [{stats['success']:2d}/50] {filename}.{fmt} -> {result.get('chunks', 0)} Chunks")
        except httpx.HTTPStatusError as e:
            stats["error"] += 1
            print(f"  [FEHLER] {filename}.{fmt}: {e.response.text[:100]}")

    # ── Zusammenfassung ──
    print("\n" + "=" * 60)
    print(" ZUSAMMENFASSUNG")
    print("=" * 60)
    print(f"  Erfolgreich hochgeladen: {stats['success']}/50")
    print(f"  Fehlgeschlagen:          {stats['error']}")
    print(f"  Chunks insgesamt:        {stats['total_chunks']}")
    print(f"  Ordner erstellt:         {len(folder_ids)}")
    print(f"  Temp-Dateien:            {OUTPUT_DIR}")
    print("=" * 60)

    if stats["error"] > 0:
        print(f"\n{stats['error']} Fehler aufgetreten. Siehe Details oben.")
        sys.exit(1)
    else:
        print("\nAlle 50 Dokumente erfolgreich generiert und hochgeladen!")


if __name__ == "__main__":
    main()
